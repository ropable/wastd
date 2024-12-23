from django.conf import settings
from django.contrib import messages
from django.utils import timezone
from django.db import connections, DatabaseError
from django.db.models import Q, Exists, OuterRef, Count, Subquery
from django.http import HttpResponseForbidden, HttpResponseRedirect, HttpResponse,HttpResponseBadRequest, JsonResponse
from django.contrib.auth.mixins import LoginRequiredMixin,UserPassesTestMixin
from django.shortcuts import render, redirect, get_object_or_404
from django.urls import reverse, reverse_lazy
from django.views import View
from django.views.generic.edit import FormMixin
from django.views.generic import TemplateView, ListView, DetailView, FormView, DeleteView
from .models import TrtPlaces, TrtSpecies, TrtLocations,TrtEntryBatchOrganisation
from django.core.paginator import Paginator
from openpyxl import Workbook
from django.contrib.auth.decorators import login_required
from django.utils.decorators import method_decorator
from django.views.decorators.http import require_GET, require_POST, require_http_methods
from django.template.loader import render_to_string
from django.core.serializers.json import DjangoJSONEncoder
import json
import csv
from datetime import timedelta
from django.core.exceptions import PermissionDenied,ValidationError
import pandas as pd
from datetime import datetime, date, time
from django.db import transaction
from django.apps import apps 
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from functools import reduce
import operator
import traceback
from django.db import IntegrityError

from wastd.utils import Breadcrumb, PaginateMixin
from .models import (
    TrtTurtles,
    TrtTags,
    TrtPitTags,
    TrtEntryBatches,
    TrtDataEntry,
    TrtPersons,
    TrtBeachPositions,
    TrtObservations,
    Template,
    TrtTagStates,
    TrtIdentification,
    TrtPitTagStatus,
    TrtTagStatus,
    TrtNestingSeason,
    TrtTissueTypes,
    TrtActivities,
    TrtIdentificationTypes,
    TrtEggCountMethods,
    TrtMeasurementTypes,
    TrtBodyParts,
    TrtDamageCodes,
    TrtYesNo,
    TrtRecordedTags,
    TrtRecordedPitTags,
    TrtMeasurements,
    TrtDamage,
    TrtCauseOfDeath,
    TrtTurtleStatus,
    TrtDocuments,
    TrtRecordedIdentification,
    TrtPitTagStates,
    TrtDatumCodes,
    TrtConditionCodes,
    TrtDamageCauseCodes
    
    
)
from .forms import TrtDataEntryForm, SearchForm, TrtEntryBatchesForm, TemplateForm, BatchesCodeForm, TrtPersonsForm, TagRegisterForm


class HomePageView(LoginRequiredMixin, TemplateView):
    """
    A view for the home page.

    Attributes:
        template_name (str): The name of the template to be used for rendering the view.
    """

    template_name = "wamtram2/home.html"


class EntryBatchesListView(LoginRequiredMixin, ListView):
    model = TrtEntryBatches
    template_name = "wamtram2/trtentrybatches_list.html"
    context_object_name = "batches"
    paginate_by = 30

    def dispatch(self, request, *args, **kwargs):
        if not (
            request.user.groups.filter(name="WAMTRAM2_VOLUNTEER").exists()
            or request.user.groups.filter(name="WAMTRAM2_TEAM_LEADER").exists()
            or request.user.groups.filter(name="WAMTRAM2_STAFF").exists()
            or request.user.is_superuser
        ):
            return HttpResponseForbidden(
                "You do not have permission to view this record"
            )
        return super().dispatch(request, *args, **kwargs)

    def get_queryset(self):
        queryset = super().get_queryset().order_by("-entry_batch_id")

        if (
            "filter" in self.request.GET
            and self.request.GET["filter"] == "no_observation_id"
        ):
            has_dataentry_no_observation_id = Exists(
                TrtDataEntry.objects.filter(
                    entry_batch_id=OuterRef("pk"), observation_id__isnull=True
                )
            )
            queryset = queryset.filter(has_dataentry_no_observation_id)

        # Use Subquery to fetch the last place_code for each batch
        last_place_code_subquery = Subquery(
            TrtDataEntry.objects.filter(entry_batch_id=OuterRef("pk"))
            .order_by("-data_entry_id")
            .values("place_code")[:1]
        )
        
        # Annotate the queryset with entry_count, last_place_code, and do_not_process_count
        queryset = queryset.annotate(
            entry_count=Count('trtdataentry'),
            last_place_code=last_place_code_subquery,
            do_not_process_count=Count(
                'trtdataentry',
                filter=Q(trtdataentry__do_not_process=True)
            )
        )

        return queryset

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)

        # Paginate the queryset
        queryset = self.get_queryset()
        paginator = Paginator(queryset, self.paginate_by)
        page_number = self.request.GET.get('page')
        page_obj = paginator.get_page(page_number)

        # Retrieve place_code from paginated object list and fetch corresponding TrtPlaces objects
        place_codes = [batch.last_place_code for batch in page_obj.object_list if batch.last_place_code]
        places = TrtPlaces.objects.filter(place_code__in=place_codes)
        places_dict = {place.place_code: place for place in places}

        # Attach TrtPlaces objects to each batch in the paginated list
        for batch in page_obj.object_list:
            batch.last_place_code_obj = places_dict.get(batch.last_place_code)
            
            batch.highlight_row = int(batch.do_not_process_count) > 0

        context['page_obj'] = page_obj
        context['is_paginated'] = page_obj.has_other_pages()
        context['object_list'] = page_obj.object_list
        context["persons"] = {
            person.person_id: person for person in TrtPersons.objects.all()
        }

        return context


class EntryBatchDetailView(LoginRequiredMixin, FormMixin, ListView):
    """
    A view for displaying list of a batch of TrtDataEntry objects.
    """

    model = TrtDataEntry
    template_name = "wamtram2/trtentrybatch_detail.html"
    context_object_name = "object_list"
    paginate_by = 30
    form_class = TrtEntryBatchesForm
    
    def get_initial(self):
        initial = super().get_initial()
        return initial

    def dispatch(self, request, *args, **kwargs):
        user = request.user
        batch_id = kwargs.get("batch_id")

        if user.is_superuser:
            return super().dispatch(request, *args, **kwargs)

        user_organisations = user.organisations.all()
        
        if not user_organisations.exists():
            raise PermissionDenied("You do not have permission to view this batch")

        org_codes = [org.code for org in user_organisations]
        has_permission = TrtEntryBatchOrganisation.objects.filter(
            trtentrybatch_id=batch_id,
            organisation__in=org_codes
        ).exists()

        if not has_permission:
            raise PermissionDenied("You do not have permission to view this batch")

        if not (
            user.groups.filter(name__in=["WAMTRAM2_VOLUNTEER", "WAMTRAM2_TEAM_LEADER", "WAMTRAM2_STAFF"]).exists()
        ):
            raise PermissionDenied("You do not have permission to view this batch")

        return super().dispatch(request, *args, **kwargs)

    def get(self, request, *args, **kwargs):
        if "batch_id" not in kwargs:
            new_batch = TrtEntryBatches.objects.create(
                pr_date_convention=False,
                entry_date=timezone.now().date()
            )
            self.kwargs["batch_id"] = new_batch.entry_batch_id
        return super().get(request, *args, **kwargs)

    def get_queryset(self):
        queryset = super().get_queryset()
        batch_id = self.kwargs.get("batch_id")
        
    
        filter_value = self.request.GET.get("filter")
        if filter_value == "needs_review":
            queryset = queryset.filter(entry_batch_id=batch_id, do_not_process=True)
        elif filter_value == "not_saved":
            queryset = queryset.filter(entry_batch_id=batch_id, observation_id__isnull=True)
        elif filter_value == "needs_review_no_message":
            queryset = queryset.filter(entry_batch_id=batch_id, do_not_process=True, error_message__isnull=True)
        else:
            queryset = queryset.filter(entry_batch_id=batch_id)
            
        return queryset.select_related('observation_id').order_by("-data_entry_id")

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        context["persons"] = {
            person.person_id: person for person in TrtPersons.objects.all()
        }

        batch = TrtEntryBatches.objects.get(entry_batch_id=self.kwargs.get("batch_id"))
        context["batch"] = batch
        initial = self.get_initial()
        context["form"] = TrtEntryBatchesForm(
            instance=batch,
            initial=initial
        )
        
        # Add a `highlight_row` attribute to each entry if it meets the conditions
        for entry in context['object_list']:
            entry.highlight_row = entry.do_not_process and entry.error_message not in ['None', 'Observation added to database']
        
        context['templates'] = Template.objects.all()

        cookies_key_prefix = self.kwargs.get("batch_id")
        context['selected_template'] = self.request.COOKIES.get(f'{cookies_key_prefix}_selected_template', '')
        # context['use_default_enterer'] = self.request.COOKIES.get(f'{cookies_key_prefix}_use_default_enterer', False)
        # context['default_enterer'] = self.request.COOKIES.get(f'{cookies_key_prefix}_default_enterer', None)

        context['cookies_key_prefix'] = cookies_key_prefix
        # context['default_enterer_value'] = context['default_enterer']
        
        # Add entries with do_not_process = True to the context
        context["do_not_process_entries"] = TrtDataEntry.objects.filter(
            entry_batch_id=batch.entry_batch_id,
            do_not_process=True
        ).order_by("-data_entry_id")
        
        # Add entries with do_not_process = False to the context
        context["process_entries"] = TrtDataEntry.objects.filter(
            entry_batch_id=batch.entry_batch_id,
            do_not_process=False
        ).order_by("-data_entry_id")
        
        return context

    def post(self, request, *args, **kwargs):
        form = self.get_form()
        form.instance.entry_batch_id = self.kwargs.get("batch_id")
        
        if form.is_valid():
            return self.form_valid(form)
        else:
            context = self.get_context_data(form=form, object_list=self.get_queryset())
            return self.render_to_response(context)

    def form_valid(self, form):
        batch = form.save(commit=False)
        entered_person_id = form.cleaned_data.get('entered_person_id')
        comments = form.cleaned_data.get('comments')
        if entered_person_id or comments:
            if not batch.entry_date:
                batch.entry_date = timezone.now()

            batch_id = batch.entry_batch_id

            existing_batch = TrtEntryBatches.objects.get(entry_batch_id=batch_id)
            batch.pr_date_convention = existing_batch.pr_date_convention
            batch.entry_date = existing_batch.entry_date
            batch.filename = existing_batch.filename
            messages.success(self.request, 'Batch detail saved')
            batch.save()
        else:
            context = self.get_context_data(form=form, object_list=self.get_queryset())
            return self.render_to_response(context)
        return HttpResponseRedirect(self.get_success_url())

    def get_success_url(self):
        batch_id = self.kwargs.get("batch_id")
        return reverse("wamtram2:entry_batch_detail", args=[batch_id])


class TrtDataEntryFormView(LoginRequiredMixin, FormView):
    """
    A form view for entering TRT data.
    """

    template_name = "wamtram2/trtdataentry_form.html"
    form_class = TrtDataEntryForm

    def dispatch(self, request, *args, **kwargs):
        # FIXME: Permission check
        if not (
            request.user.groups.filter(name="WAMTRAM2_VOLUNTEER").exists()
            or request.user.groups.filter(name="WAMTRAM2_TEAM_LEADER").exists()
            or request.user.groups.filter(name="WAMTRAM2_STAFF").exists()
            or request.user.is_superuser
        ):
            return HttpResponseForbidden(
                "You do not have permission to view this record"
            )
        return super().dispatch(request, *args, **kwargs)

    def get_form_kwargs(self):
        """
        Returns the keyword arguments for instantiating the form.

        If an entry_id is provided in the URL, retrieves the corresponding TrtDataEntry instance
        and adds it as the 'instance' argument in the form kwargs. If no entry_id is provided,
        a new, blank form is instantiated.

        Returns:
            dict: The keyword arguments for instantiating the form.
        """
        kwargs = super().get_form_kwargs()
        entry_id = self.kwargs.get("entry_id")
        if entry_id:
            entry = get_object_or_404(TrtDataEntry.objects.select_related('turtle_id'), data_entry_id=entry_id)
            kwargs["instance"] = entry

        return kwargs

    def get_template_data(self, template_key):
        try:
            template = Template.objects.get(pk=template_key)
            return {
                'place_code': template.place_code,
                'species_code': template.species_code,
                'sex': template.sex,
                'location_code': template.location_code
            }
        except Template.DoesNotExist:
            return None

    def get_initial(self):
        initial = super().get_initial()
        batch_id = self.kwargs.get("batch_id")
        turtle_id = self.kwargs.get("turtle_id")
        entry_id = self.kwargs.get("entry_id")
        cookies_key_prefix = batch_id
        
        do_not_process = self.request.COOKIES.get(f'{cookies_key_prefix}_do_not_process') == 'true'
        if do_not_process:
            initial['do_not_process'] = True
            initial['comments'] = "The data in the sheet doesn't match the database."
    
        
        tag_id = self.request.COOKIES.get(f'{cookies_key_prefix}_tag_id')
        tag_type = self.request.COOKIES.get(f'{cookies_key_prefix}_tag_type')
        tag_side = self.request.COOKIES.get(f'{cookies_key_prefix}_tag_side')

        selected_template = self.request.COOKIES.get(f'{cookies_key_prefix}_selected_template')
        
        # If a tag is selected, populate the form with the tag data
        if tag_id and tag_type:
            if tag_type == 'recapture_tag':
                if tag_side == 'L':
                    initial['recapture_left_tag_id'] = tag_id
                elif tag_side == 'R':
                    initial['recapture_right_tag_id'] = tag_id
            elif tag_type == 'recapture_pit_tag':
                initial['recapture_pittag_id'] = tag_id
                
        if batch_id:
            try:
                batch = TrtEntryBatches.objects.get(entry_batch_id=batch_id)
                if batch.template:
                    selected_template = str(batch.template.template_id)
            except TrtEntryBatches.DoesNotExist:
                pass

        if selected_template:
            template_data = self.get_template_data(selected_template)
            if template_data:
                place_code = template_data.get('place_code') or ""
                initial['default_place_code'] = place_code or ""
                self.default_place_code = place_code
                default_place_obj = TrtPlaces.objects.filter(place_code=self.default_place_code).first()
                if default_place_obj:
                    self.default_place_full_name = default_place_obj.get_full_name()
                    initial['default_place_full_name'] = self.default_place_full_name or ""
                if not turtle_id:
                    initial['species_code'] = template_data.get('species_code') or ""
                    initial['sex'] = template_data.get('sex') or ""
                

        if batch_id:
            initial["entry_batch"] = get_object_or_404(TrtEntryBatches, entry_batch_id=batch_id)

        if turtle_id:
            turtle = get_object_or_404(TrtTurtles.objects.prefetch_related('trttags_set', 'trtpittags_set'), pk=turtle_id)
            initial["turtle_id"] = turtle_id
            initial["species_code"] = turtle.species_code
            initial["sex"] = turtle.sex

        if entry_id:
            trtdataentry = get_object_or_404(TrtDataEntry, data_entry_id=entry_id)

            if trtdataentry.observation_date:
                adjusted_date = trtdataentry.observation_date - timedelta(hours=8)
                initial['observation_date'] = adjusted_date
            
            measured_by = trtdataentry.measured_by
            recorded_by = trtdataentry.recorded_by
            tagged_by = trtdataentry.tagged_by
            entered_by = trtdataentry.entered_by
            place_code = trtdataentry.place_code
            
            if place_code:
                place = TrtPlaces.objects.filter(place_code=place_code).first()
                if place:
                    initial["place_code"] = place_code
                    self.place_full_name = place.get_full_name()
                else:
                    self.place_full_name = ""
            
            if measured_by:
                parts = measured_by.split(" ", 1)
                first_name, last_name = parts
                person = TrtPersons.objects.filter(first_name=first_name, surname=last_name).first()
                if person:
                    initial["measured_by_id"] = person.person_id
                    self.measured_by_full_name = measured_by
                else:
                    self.measured_by_full_name = ""

            if recorded_by:
                parts = recorded_by.split(" ", 1)
                first_name, last_name = parts
                person = TrtPersons.objects.filter(first_name=first_name, surname=last_name).first()
                if person:
                    initial["recorded_by_id"] = person.person_id
                    self.recorded_by_full_name = recorded_by
                else:
                    self.recorded_by_full_name = ""

            if tagged_by:
                parts = tagged_by.split(" ", 1)
                first_name, last_name = parts
                person = TrtPersons.objects.filter(first_name=first_name, surname=last_name).first()
                if person:
                    initial["tagged_by_id"] = person.person_id
                    self.tagged_by_full_name = tagged_by
                else:
                    self.tagged_by_full_name = ""

            if entered_by:
                parts = entered_by.split(" ", 1)
                first_name, last_name = parts
                person = TrtPersons.objects.filter(first_name=first_name, surname=last_name).first()
                if person:
                    initial["entered_by_id"] = person.person_id
                    self.entered_by_full_name = entered_by
                else:
                    self.entered_by_full_name = ""
        return initial

    def form_valid(self, form):
        batch_id = form.cleaned_data["entry_batch"].entry_batch_id
        do_not_process_cookie_name = f"{batch_id}_do_not_process"
        do_not_process_cookie_value = self.request.COOKIES.get(do_not_process_cookie_name)
        if do_not_process_cookie_value == 'true':
            form.instance.do_not_process = True
        entry = form.save()
        #success_url = reverse("wamtram2:find_turtle", args=[batch_id])
        success_url = FindTurtleView.get_clear_cookies_url(batch_id)
        
        if form.instance.do_not_process:
            message = f"Entry created successfully and will be reviewed later. Please write the Entry ID: {entry.data_entry_id} on the data sheet"
            message_tag = 'warning'
        else:
            message = f"Entry created successfully. Entry ID: {entry.data_entry_id}"
            message_tag = 'success'
        
        messages.add_message(self.request, getattr(messages, message_tag.upper()), message)
        
        if self.request.headers.get('x-requested-with') == 'XMLHttpRequest':
            return JsonResponse({
                'success': True, 
                'redirect_url': success_url,
                'message': message,
                'message_tag': message_tag
            })
        else:
            return redirect(success_url)


    def form_invalid(self, form):
        error_message = "Error saving the entry. If you cannot resolve the issue, please set aside this data sheet for admin review and continue with the next data sheet."
        
        messages.error(self.request, error_message)
        
        if self.request.headers.get('x-requested-with') == 'XMLHttpRequest':
            return JsonResponse({
                'success': False, 
                'errors': form.errors,
                'message': error_message
            })
        else:
            return self.render_to_response(self.get_context_data(form=form))

    def get_context_data(self, **kwargs):
        """
        Returns the context data for rendering the template.

        Adds the entry_id and entry objects to the context.

        Returns:
            dict: The context data.
        """
        context = super().get_context_data(**kwargs)
        entry_id = self.kwargs.get("entry_id")
        batch_id = self.kwargs.get("batch_id")
        cookies_key_prefix = batch_id
        form = context.get('form')
        
        if form.is_bound:
            context.update({
                'entered_by_id': form.data.get('entered_by_id'),
                'recorded_by_id': form.data.get('recorded_by_id'),
                'measured_by_id': form.data.get('measured_by_id'),
                'tagged_by_id': form.data.get('tagged_by_id'),
                'place_code': form.data.get('place_code'),
                'entered_by_name': self.get_person_name(form.data.get('entered_by_id')),
                'recorded_by_name': self.get_person_name(form.data.get('recorded_by_id')),
                'measured_by_name': self.get_person_name(form.data.get('measured_by_id')),
                'tagged_by_name': self.get_person_name(form.data.get('tagged_by_id')),
                'place_name': self.get_place_name(form.data.get('place_code')),
            })

        if entry_id:
            entry = get_object_or_404(TrtDataEntry.objects.select_related('turtle_id'), data_entry_id=entry_id)
            context["entry_id"] = entry_id  # Editing existing entry
            context["entry"] = entry
            
            if entry.observation_id:
                context["observation"] = entry.observation_id
            else:
                context["observation"] = None
        
        if batch_id:
            context["batch_id"] = batch_id  # Creating new entry in batch
            batch = TrtEntryBatches.objects.get(entry_batch_id=batch_id)
            if batch.template:
                context["selected_template"] = str(batch.template.template_id)
            else:
                context["selected_template"] = self.request.COOKIES.get(f'{cookies_key_prefix}_selected_template') or None
            # context["use_default_enterer"] = self.request.COOKIES.get(f'{cookies_key_prefix}_use_default_enterer', False)
            # context["default_enterer"] = self.request.COOKIES.get(f'{cookies_key_prefix}_default_enterer', '')
            # Add the tag id and tag type to the context data
            context["cookie_tag_id"] = self.request.COOKIES.get(f'{cookies_key_prefix}_tag_id')
            context["cookie_tag_type"] = self.request.COOKIES.get(f'{cookies_key_prefix}_tag_type')
            context["cookie_tag_side"] = self.request.COOKIES.get(f'{cookies_key_prefix}_tag_side')
            context["unprocessed_turtle"] = bool(context["cookie_tag_id"])

            context["default_enterer_full_name"] = getattr(self, 'default_enterer_full_name', '')
            context["default_place_full_name"] = getattr(self, 'default_place_full_name', '')
            context["default_place_code"] = getattr(self, 'default_place_code', '')
        
            context['measured_by_full_name'] = getattr(self, 'measured_by_full_name', '')
            context['recorded_by_full_name'] = getattr(self, 'recorded_by_full_name', '')
            context['tagged_by_full_name'] = getattr(self, 'tagged_by_full_name', '')
            context['entered_by_full_name'] = getattr(self, 'entered_by_full_name', '')
            context['place_name'] = getattr(self, 'place_full_name', '')

        return context

    def get_person_name(self, person_id):
        if person_id:
            try:
                person = TrtPersons.objects.get(person_id=person_id)
                return f"{person.first_name} {person.surname}"
            except TrtPersons.DoesNotExist:
                return ""
        return ""
        
    def get_place_name(self, place_code):
        if place_code:
            try:
                place = TrtPlaces.objects.get(place_code=place_code)
                return place.get_full_name()
            except TrtPlaces.DoesNotExist:
                return ""
        return ""


class DeleteBatchView(LoginRequiredMixin, View):
    def dispatch(self, request, *args, **kwargs):
        # FIXME: Permission check
        if not (
            request.user.groups.filter(name="WAMTRAM2_STAFF").exists()
            or request.user.is_superuser
        ):
            return HttpResponseForbidden(
                "You do not have permission to view this record"
            )
        return super().dispatch(request, *args, **kwargs)

    def post(self, request, batch_id):
        batch = get_object_or_404(TrtEntryBatches, entry_batch_id=batch_id)
        batch.delete()
        return redirect("wamtram2:batches_curation")


class ValidateDataEntryBatchView(LoginRequiredMixin, View):
    """
    View class for validating a data entry batch.

    This view executes a stored procedure to validate the data in a batch
    identified by the 'batch_id' parameter. If the validation is successful,
    a success message is added to the request's messages framework. If there
    is a database error, an error message is added instead.

    After the validation, the view redirects to the 'entry_batch_detail' view
    with the 'batch_id' parameter.

    Attributes:
        - request: The HTTP request object.
        - args: Additional positional arguments passed to the view.
        - kwargs: Additional keyword arguments passed to the view.
    """

    def dispatch(self, request, *args, **kwargs):
        # FIXME: Permission check
        if not (
            request.user.groups.filter(name="WAMTRAM2_TEAM_LEADER").exists()
            or request.user.groups.filter(name="WAMTRAM2_STAFF").exists()
            or request.user.is_superuser
        ):
            return HttpResponseForbidden(
                "You do not have permission to view this record"
            )
        return super().dispatch(request, *args, **kwargs)

    def get(self, request, *args, **kwargs):
        try:
            with connections["wamtram2"].cursor() as cursor:
                cursor.execute(
                    "EXEC dbo.ValidateDataEntryBatchWEB @ENTRY_BATCH_ID = %s",
                    [self.kwargs["batch_id"]],
                )
                messages.add_message(request, messages.INFO, "Validation finished.")
        except DatabaseError as e:
            messages.add_message(
                request, messages.ERROR, "Database error: {}".format(e)
            )
            
        return_to = request.GET.get('return_to')
        
        if return_to == 'curation':
            return redirect("wamtram2:entries_curation", batch_id=self.kwargs["batch_id"])
        else:
            return redirect("wamtram2:entry_batch_detail", batch_id=self.kwargs["batch_id"])


class DeleteEntryView(LoginRequiredMixin,DeleteView):
    model = TrtDataEntry
    success_url = reverse_lazy('wamtram2:batches_curation')
    
    def dispatch(self, request, *args, **kwargs):
        # FIXME: Permission check
        if not (
            request.user.groups.filter(name="WAMTRAM2_STAFF").exists()
            or request.user.is_superuser
        ):
            return HttpResponseForbidden(
                "You do not have permission to view this record"
            )
        return super().dispatch(request, *args, **kwargs)

    def get_success_url(self):
        batch_id = self.kwargs['batch_id']
        return reverse_lazy('wamtram2:entry_batch_detail', kwargs={'batch_id': batch_id})


class ProcessDataEntryBatchView(LoginRequiredMixin, View):
    """
    View class for processing a data entry batch.

    This view executes a stored procedure to process a data entry batch
    identified by the batch ID provided in the URL parameters. It uses the
    'wamtram2' database connection and redirects the user to the detail page
    of the processed batch.

    Attributes:
        None

    Methods:
        get: Handles the GET request and executes the stored procedure.

    Raises:
        DatabaseError: If there is an error executing the stored procedure.

    Returns:
        HttpResponseRedirect: Redirects the user to the detail page of the
        processed batch.
    """

    def dispatch(self, request, *args, **kwargs):
        # FIXME: Permission check
        if not (
            request.user.groups.filter(name="WAMTRAM2_STAFF").exists()
            or request.user.is_superuser
        ):
            return HttpResponseForbidden(
                "You do not have permission to view this record"
            )
        return super().dispatch(request, *args, **kwargs)

    def get(self, request, *args, **kwargs):
        try:
            with connections["wamtram2"].cursor() as cursor:
                cursor.execute(
                    "EXEC dbo.EntryBatchProcessWEB @ENTRY_BATCH_ID = %s;",
                    [self.kwargs["batch_id"]],
                )
                messages.add_message(request, messages.INFO, "Processing finished.")
        except DatabaseError as e:
            messages.add_message(
                request, messages.ERROR, "Database error: {}".format(e)
            )
            
        return_to = request.GET.get('return_to')
        
        if return_to == 'curation':
            return redirect("wamtram2:entries_curation", batch_id=self.kwargs["batch_id"])
        else:
            return redirect("wamtram2:entry_batch_detail", batch_id=self.kwargs["batch_id"])


class FindTurtleView(LoginRequiredMixin, View):
    """
    View class for finding a turtle based on tag and pit tag ID.
    """
    template_name= "wamtram2/find_turtle.html"
    def dispatch(self, request, *args, **kwargs):
        if not (
            request.user.groups.filter(name="WAMTRAM2_VOLUNTEER").exists()
            or request.user.groups.filter(name="WAMTRAM2_TEAM_LEADER").exists()
            or request.user.groups.filter(name="WAMTRAM2_STAFF").exists()
            or request.user.is_superuser
        ):
            return HttpResponseForbidden("You do not have permission to view this record")
        return super().dispatch(request, *args, **kwargs)

    def get(self, request, *args, **kwargs):
        batch_id = kwargs.get("batch_id")
        form = SearchForm(initial={"batch_id": batch_id})
        
        clear_cookies = request.GET.get('clear_cookies', 'false') == 'true'
    
        if clear_cookies:
            response = render(request, "wamtram2/find_turtle.html", {
                "form": form,
                "batch_id": batch_id,
                "batch": TrtEntryBatches.objects.filter(entry_batch_id=batch_id).first(),
                "template_name": "No template associated",
            })
            self.clear_search_cookies(response, batch_id)
            return response
            
        
        no_turtle_found = request.COOKIES.get(f'{batch_id}_no_turtle_found') == "true"
        tag_id = request.COOKIES.get(f'{batch_id}_tag_id')
        tag_type = request.COOKIES.get(f'{batch_id}_tag_type')
        tag_side = request.COOKIES.get(f'{batch_id}_tag_side')
        turtle = None
        first_observation_date = None
        latest_site = None
        batch = None
        template_name = "No template associated"
        new_tag_entry = None
        if tag_id and tag_type and not turtle:
            new_tag_entry = TrtDataEntry.objects.filter(
                Q(new_left_tag_id__tag_id=tag_id) |
                Q(new_left_tag_id_2__tag_id=tag_id) |
                Q(new_right_tag_id__tag_id=tag_id) |
                Q(new_right_tag_id_2__tag_id=tag_id) |
                Q(new_pittag_id__pittag_id=tag_id) |
                Q(new_pittag_id_2__pittag_id=tag_id) |
                Q(new_pittag_id_3__pittag_id=tag_id) |
                Q(new_pittag_id_4__pittag_id=tag_id),
                observation_id__isnull=True,
            ).select_related('entry_batch', 'place_code', 'species_code').order_by('-entry_batch__entry_date').first()

        if batch_id:
            batch = TrtEntryBatches.objects.filter(entry_batch_id=batch_id).first()
            if batch and batch.template:
                template_name = batch.template.name

        if tag_id and tag_type and not no_turtle_found:
            tag = TrtTags.objects.select_related('turtle').filter(tag_id=tag_id).first()
            if tag:
                turtle = tag.turtle
            else:
                pit_tag = TrtPitTags.objects.select_related('turtle').filter(pittag_id=tag_id).first()
                if pit_tag:
                    turtle = pit_tag.turtle

            if turtle:
                first_observation = turtle.trtobservations_set.order_by('observation_date').first()
                if first_observation:
                    first_observation_date = first_observation.observation_date

                latest_observation = turtle.trtobservations_set.order_by('-observation_date').first()
                if latest_observation and latest_observation.place_code:
                    latest_site = latest_observation.place_code.place_name

        return render(request, "wamtram2/find_turtle.html", {
            "form": form,
            "turtle": turtle,
            "no_turtle_found": no_turtle_found,
            "tag_id": tag_id,
            "tag_type": tag_type,
            "tag_side": tag_side,
            "first_observation_date": first_observation_date,
            "latest_site": latest_site,
            "batch_id": batch_id,
            "batch": batch,
            "template_name": template_name,
            "new_tag_entry": new_tag_entry,
        })
        
    def clear_search_cookies(self, response, batch_id):
        cookies_to_clear = [
            f'{batch_id}_tag_id',
            f'{batch_id}_tag_type',
            f'{batch_id}_tag_side',
            f'{batch_id}_no_turtle_found',
            f'{batch_id}_do_not_process'
        ]
        for cookie_name in cookies_to_clear:
            response.delete_cookie(cookie_name)

    @staticmethod
    def get_clear_cookies_url(batch_id):
        return reverse("wamtram2:find_turtle", args=[batch_id]) + '?clear_cookies=true'
    
    def set_cookie(self, response, batch_id, tag_id=None, tag_type=None, tag_side=None, no_turtle_found=False, do_not_process=False):
        if tag_id:
            response.set_cookie(f'{batch_id}_tag_id', tag_id, max_age=63072000)
        if tag_type:
            response.set_cookie(f'{batch_id}_tag_type', tag_type, max_age=63072000)
        if tag_side:
            response.set_cookie(f'{batch_id}_tag_side', tag_side, max_age=63072000)
        response.set_cookie(f'{batch_id}_no_turtle_found', 'true' if no_turtle_found else 'false', max_age=63072000)
        response.set_cookie(f'{batch_id}_do_not_process', 'true' if do_not_process else 'false', max_age=63072000)
        return response

    def post(self, request, *args, **kwargs):
        batch_id = kwargs.get("batch_id")
        form = SearchForm(request.POST, initial={"batch_id": batch_id})
        no_turtle_found = False
        tag_type = None
        tag_id = None
        tag_side = None
        turtle = None
        create_and_review = request.POST.get('create_and_review') == 'true'
        new_tag_entry = None
        batch = None
        template_name = "No template associated"
        existing_turtle_entry = None
        
        if batch_id:
            batch = TrtEntryBatches.objects.filter(entry_batch_id=batch_id).first()
            if batch and batch.template:
                template_name = batch.template.name

        if form.is_valid():
            tag_id = form.cleaned_data["tag_id"].upper()

            if not create_and_review:
                tag = TrtTags.objects.select_related('turtle').filter(tag_id=tag_id).first()
                if tag:
                    turtle = tag.turtle
                    tag_type = "recapture_tag"
                    tag_side = tag.side
                else:
                    pit_tag = TrtPitTags.objects.select_related('turtle').filter(pittag_id=tag_id).first()
                    if pit_tag:
                        turtle = pit_tag.turtle
                        tag_type = "recapture_pit_tag"
                    else:
                        tag_type = "unknown_tag"
                        
                if not turtle:
                    existing_turtle_entry = TrtDataEntry.objects.filter(
                        Q(new_left_tag_id__tag_id=tag_id) |
                        Q(new_left_tag_id_2__tag_id=tag_id) |
                        Q(new_right_tag_id__tag_id=tag_id) |
                        Q(new_right_tag_id_2__tag_id=tag_id) |
                        Q(new_pittag_id__pittag_id=tag_id) |
                        Q(new_pittag_id_2__pittag_id=tag_id) |
                        Q(new_pittag_id_3__pittag_id=tag_id) |
                        Q(new_pittag_id_4__pittag_id=tag_id),
                        turtle_id__isnull=False,  
                        observation_id__isnull=True
                    ).select_related(
                        'turtle_id', 
                        'entry_batch', 
                        'place_code', 
                        'species_code'
                    ).order_by('-entry_batch__entry_date').first()
                    
                    if existing_turtle_entry:
                        turtle = existing_turtle_entry.turtle_id
                        if any([
                            str(existing_turtle_entry.new_left_tag_id).upper() == str(tag_id).upper(),
                            str(existing_turtle_entry.new_left_tag_id_2).upper() == str(tag_id).upper()
                        ]):
                            tag_type = "recapture_tag"
                            tag_side = "L"
                        elif any([
                            str(existing_turtle_entry.new_right_tag_id).upper() == str(tag_id).upper(),
                            str(existing_turtle_entry.new_right_tag_id_2).upper() == str(tag_id).upper()
                        ]):
                            tag_type = "recapture_tag"
                            tag_side = "R"
                        else:
                            tag_type = "recapture_pit_tag"
                            tag_side = None
                    else:
                        new_tag_entry = TrtDataEntry.objects.filter(
                            Q(new_left_tag_id__tag_id=tag_id) |
                            Q(new_left_tag_id_2__tag_id=tag_id) |
                            Q(new_right_tag_id__tag_id=tag_id) |
                            Q(new_right_tag_id_2__tag_id=tag_id) |
                            Q(new_pittag_id__pittag_id=tag_id) |
                            Q(new_pittag_id_2__pittag_id=tag_id) |
                            Q(new_pittag_id_3__pittag_id=tag_id) |
                            Q(new_pittag_id_4__pittag_id=tag_id),
                            observation_id__isnull=True,
                            turtle_id__isnull=True
                        ).select_related(
                            'entry_batch', 
                            'place_code', 
                            'species_code'
                        ).order_by('-entry_batch__entry_date').first()

                        if new_tag_entry:
                            if any([
                                str(new_tag_entry.new_left_tag_id).upper() == str(tag_id).upper(),
                                str(new_tag_entry.new_left_tag_id_2).upper() == str(tag_id).upper()
                            ]):
                                tag_type = "recapture_tag"
                                tag_side = "L"
                            elif any([
                                str(new_tag_entry.new_right_tag_id).upper() == str(tag_id).upper(),
                                str(new_tag_entry.new_right_tag_id_2).upper() == str(tag_id).upper()
                            ]):
                                tag_type = "recapture_tag"
                                tag_side = "R"
                            else:
                                tag_type = "recapture_pit_tag"
                                tag_side = None
                        else:
                            no_turtle_found = True

                if turtle:
                    if existing_turtle_entry:
                        response = render(request, "wamtram2/find_turtle.html", {
                            "form": form,
                            "turtle": turtle,
                            "existing_turtle_entry": existing_turtle_entry,
                            "tag_id": tag_id,
                            "tag_type": tag_type,
                            "tag_side": tag_side,
                            "batch_id": batch_id,
                            "batch": batch,
                            "template_name": template_name,
                        })
                    else:
                        response = redirect(reverse('wamtram2:find_turtle', kwargs={'batch_id': batch_id}))
                    return self.set_cookie(response, batch_id, tag_id, tag_type, tag_side)
                                
                elif new_tag_entry:
                    response = render(request, "wamtram2/find_turtle.html", {
                        "form": form,
                        "turtle": turtle,
                        "new_tag_entry": new_tag_entry,
                        "no_turtle_found": no_turtle_found,
                        "tag_id": tag_id,
                        "tag_type": tag_type,
                        "tag_side": tag_side,
                        "batch_id": batch_id,
                        "batch": batch,
                        "template_name": template_name,
                        "unprocessed_turtle": True
                    })
                    return self.set_cookie(response, batch_id, tag_id, tag_type, tag_side, no_turtle_found)
                else:
                    response = redirect(reverse('wamtram2:find_turtle', kwargs={'batch_id': batch_id}))
                    return self.set_cookie(response, batch_id, tag_id, tag_type, tag_side, no_turtle_found=True)
            else:
                tag_type = request.POST.get('tag_type', 'unknown_tag')
                tag_side = request.POST.get('tag_side', None)
                response = redirect(reverse('wamtram2:newtrtdataentry', kwargs={'batch_id': batch_id}))
                return self.set_cookie(response, batch_id, tag_id, tag_type, tag_side, do_not_process=True)
        else:
            response = render(request, "wamtram2/find_turtle.html", {
                "form": form,
                "no_turtle_found": no_turtle_found,
                "tag_id": tag_id,
                "tag_type": tag_type,
                "tag_side": tag_side,
                "batch_id": batch_id,
            })

        return self.set_cookie(response, batch_id, tag_id, tag_type, tag_side)
    
    
class ObservationDetailView(LoginRequiredMixin, DetailView):
    model = TrtObservations
    template_name = "wamtram2/observation_detail.html"
    
    def dispatch(self, request, *args, **kwargs):
        if not (
            request.user.groups.filter(name="WAMTRAM2_VOLUNTEER").exists()
            or request.user.groups.filter(name="WAMTRAM2_TEAM_LEADER").exists()
            or request.user.groups.filter(name="WAMTRAM2_STAFF").exists()
            or request.user.is_superuser
        ):
            return HttpResponseForbidden("You do not have permission to view this record")
        return super().dispatch(request, *args, **kwargs)


    def get_object(self, queryset=None):
        if queryset is None:
            queryset = self.get_queryset()
        return super().get_object(queryset)

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        obj = self.object
        
        context["observation"] = obj
        
        context["tags"] = obj.trtrecordedtags_set.all()
        context["pittags"] = obj.trtrecordedpittags_set.all()
        context["measurements"] = obj.trtmeasurements_set.all()
        
        if obj.place_code:
            context["place_full_name"] = obj.place_code.get_full_name()
        else:
            context["place_full_name"] = ""
        return context


class TurtleListView(LoginRequiredMixin, PaginateMixin, ListView):
    """
    View class for displaying a list of turtles.

    Attributes:
        model (Model): The model class representing the turtles.
        paginate_by (int): The number of turtles to display per page.
    """

    model = TrtTurtles
    paginate_by = 50
    template_name = "wamtram2/trtturtles_list.html"

    def get_context_data(self, **kwargs):
        """
        Retrieves the context data for rendering the template.

        Returns:
            dict: The context data.
        """
        context = super().get_context_data(**kwargs)
        context["page_title"] = f"{settings.SITE_CODE} | WAMTRAM2"
        # Pass in any query string
        if "q" in self.request.GET:
            context["query_string"] = self.request.GET["q"]
        context["breadcrumbs"] = (
            Breadcrumb("Home", reverse("home")),
            Breadcrumb("Tagged turtles", None),
        )
        return context

    def get_queryset(self):
        """
        Retrieves the queryset of turtles to be displayed.

        Returns:
            QuerySet: The queryset of turtles.
        """
        qs = super().get_queryset()
        # General-purpose search uses the `q` parameter.
        if "q" in self.request.GET and self.request.GET["q"]:
            q = self.request.GET["q"]
            qs = qs.filter(
                Q(pk__icontains=q)
                | Q(trttags__tag_id__icontains=q)
                | Q(trtpittags__pittag_id__icontains=q)
            ).distinct()

        return qs.order_by("pk")


class TurtleDetailView(LoginRequiredMixin, DetailView):
    """
    View class for displaying and exporting the details of a turtle.

    Attributes:
        model (Model): The model class representing the turtle.
        template_name (str): The template used for displaying turtle details.
    """

    model = TrtTurtles
    template_name = "wamtram2/trtturtles_detail.html"
    
    def dispatch(self, request, *args, **kwargs):
        """
        Check user permissions before processing the request.
        Only allows access to authorized users.
        """
        if not (
            request.user.groups.filter(name="WAMTRAM2_VOLUNTEER").exists()
            or request.user.groups.filter(name="WAMTRAM2_TEAM_LEADER").exists()
            or request.user.groups.filter(name="WAMTRAM2_STAFF").exists()
            or request.user.is_superuser
        ):
            return HttpResponseForbidden("You do not have permission to view this record")
        return super().dispatch(request, *args, **kwargs)
    
    def get(self, request, *args, **kwargs):
        """
        Handle GET requests, either display detail view or export Word document
        """
        if 'export' in request.path:
            return self.export_word(request, *args, **kwargs)
        return super().get(request, *args, **kwargs)

    def get_context_data(self, **kwargs):
        """
        Retrieve and prepare the context data for the template.

        Returns:
            dict: Context data including turtle details, tags, observations, and samples.
        """
        context = super().get_context_data(**kwargs)
        obj = self.get_object()
        
        # Get unique PIT tags
        pittags = obj.recorded_pittags.all().order_by('pittag_id', '-observation_id')
        seen = set()
        unique_pittags = []
        for tag in pittags:
            if tag.pittag_id_id not in seen:
                unique_pittags.append(tag)
                seen.add(tag.pittag_id_id)
        
        # Get observations and measurements
        observations = obj.trtobservations_set.all()
        observations_data = []
        for obs in observations:
            obs_data = {
                'observation': obs,
                'measurements': obs.trtmeasurements_set.all()
            }
            observations_data.append(obs_data)
            
        identifications = TrtIdentification.objects.filter(turtle_id=obj.pk)
        
        context.update({
            "page_title": f"{settings.SITE_CODE} | WAMTRAM2 | {obj.pk}",
            "tags": obj.trttags_set.all(),
            "pittags": unique_pittags,
            "observations_data": observations_data,
            "samples": obj.trtsamples_set.all(),
            "identifications": identifications 
        })
                
        return context

    def export_word(self, request, *args, **kwargs):
        """
        Export turtle information to a Word document.
        
        Returns:
            HttpResponse: Word document as a downloadable file.
        """
        turtle = self.get_object()
        
        # Create new document
        doc = Document()
        
        section = doc.sections[0]
        header = section.header
        header_para = header.paragraphs[0]
        header_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = header_para.add_run()
        run.add_picture('wastd/static/android-chrome-192x192.png', width=Inches(0.4))
    
        # Title with formatting
        title_para = doc.add_paragraph()
        title_run = title_para.add_run('W.A. Marine Turtles Conservation Database - Turtle Information Sheet')
        title_run.font.size = Pt(18)
        title_run.font.color.rgb = RGBColor(31,73,125)
        title_run.font.bold = True
        title_para.space_after = Pt(12)
        
        # Basic information
        doc.add_paragraph(f'Turtle ID: {turtle.pk}')
        doc.add_paragraph(f'Species: {turtle.species_code or ""}')
        doc.add_paragraph(f'Sex: {turtle.sex or "Unknown"}')
        doc.add_paragraph(f'Status: {turtle.turtle_status or ""}')
        doc.add_paragraph(f'Cause of Death: {turtle.cause_of_death or ""}')

        def add_section_title(text):
            para = doc.add_paragraph()
            run = para.add_run(text)
            run.font.size = Pt(16)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0,32,96)
            para.space_before = Pt(12)
            para.space_after = Pt(6)
            return para
        
        # Tags information
        add_section_title('Tags Information:')
        flipper_tags = TrtTags.objects.filter(turtle=turtle)
        pit_tags = TrtPitTags.objects.filter(turtle=turtle)
        
        if flipper_tags.exists() or pit_tags.exists():
            table = doc.add_table(rows=1, cols=4)
            table.style = 'Table Grid'
            header_cells = table.rows[0].cells
            header_cells[0].text = 'Tag Type'
            header_cells[1].text = 'Tag ID'
            header_cells[2].text = 'Status'
            header_cells[3].text = 'Comments'
            
            for tag in flipper_tags:
                row_cells = table.add_row().cells
                row_cells[0].text = 'Flipper'
                row_cells[1].text = str(tag.tag_id)
                row_cells[2].text = str(tag.tag_status) if tag.tag_status else ''
                row_cells[3].text = str(tag.comments or '')
                
            for tag in pit_tags:
                row_cells = table.add_row().cells
                row_cells[0].text = 'PIT'
                row_cells[1].text = str(tag.pittag_id)
                row_cells[2].text = str(tag.pit_tag_status) if tag.pit_tag_status else ''
                row_cells[3].text = str(tag.comments or '')
        else:
            doc.add_paragraph('No tags recorded')
        
        # Other identification history
        add_section_title('Other Identification History:')
        identifications = TrtIdentification.objects.filter(turtle_id=turtle.pk)
        if identifications.exists():
            table = doc.add_table(rows=1, cols=3)
            table.style = 'Table Grid'
            header_cells = table.rows[0].cells
            header_cells[0].text = 'Identification Type'
            header_cells[1].text = 'Identifier'
            header_cells[2].text = 'Comments'
            
            for ident in identifications:
                row_cells = table.add_row().cells
                row_cells[0].text = str(ident.identification_type)
                row_cells[1].text = str(ident.identifier)
                row_cells[2].text = str(ident.comments or '')
        else:
            doc.add_paragraph('No identification history recorded')
        
        # Observations
        add_section_title('Observations:')
        observations = turtle.trtobservations_set.all()
        if observations:
            table = doc.add_table(rows=1, cols=3)
            table.style = 'Table Grid'
            header_cells = table.rows[0].cells
            header_cells[0].text = 'Date'
            header_cells[1].text = 'Place'
            header_cells[2].text = 'Activity'
            
            for obs in observations:
                row_cells = table.add_row().cells
                row_cells[0].text = obs.observation_date.strftime('%d/%m/%Y %H:%M:%S')
                row_cells[1].text = str(obs.place_code.get_full_name() if obs.place_code else '')
                row_cells[2].text = str(obs.activity_code if obs.activity_code else '')
        else:
            doc.add_paragraph('No observations recorded')
            
        # All Measurements in one table
        add_section_title('Measurements:')
        all_measurements = []
        for obs in observations:
            measurements = obs.trtmeasurements_set.all()
            all_measurements.extend(measurements)
            
        if all_measurements:
            table = doc.add_table(rows=1, cols=4)
            table.style = 'Table Grid'
            header_cells = table.rows[0].cells
            header_cells[0].text = 'Date'
            header_cells[1].text = 'Measurement Type'
            header_cells[2].text = 'Value'
            header_cells[3].text = 'Comments'
            
            for m in all_measurements:
                row_cells = table.add_row().cells
                row_cells[0].text = m.observation.observation_date.strftime('%d/%m/%Y %H:%M:%S')
                row_cells[1].text = str(m.measurement_type)
                row_cells[2].text = str(m.measurement_value)
                row_cells[3].text = str(m.comments or '')
        else:
            doc.add_paragraph('No measurements recorded')
        
        # Samples
        add_section_title('Samples:')
        samples = turtle.trtsamples_set.all()
        if samples:
            table = doc.add_table(rows=1, cols=4)
            table.style = 'Table Grid'
            header_cells = table.rows[0].cells
            header_cells[0].text = 'Tissue'
            header_cells[1].text = 'Date'
            header_cells[2].text = 'Label'
            header_cells[3].text = 'Comments'
            
            for sample in samples:
                row_cells = table.add_row().cells
                row_cells[0].text = str(sample.tissue_type)
                row_cells[1].text = sample.sample_date.strftime('%d/%m/%Y') if sample.sample_date else ''
                row_cells[2].text = str(sample.sample_label or '')
                row_cells[3].text = str(sample.comments or '')
        else:
            doc.add_paragraph('No samples recorded')
        # Add footer
        doc.add_paragraph(f'WAMTRAM - {timezone.now().strftime("%d-%b-%Y")} copy. Department of Biodiversity, Conservation and Attractions')
        
        # Prepare response
        response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        response['Content-Disposition'] = f'attachment; filename=turtle_{turtle.pk}_report.docx'
        
        doc.save(response)
        return response
    
    
SEX_CHOICES = [
    ("F", "Female"),
    ("M", "Male"),
    ("I", "Indeterminate"),
]
class TemplateManageView(LoginRequiredMixin, FormView):
    template_name = 'wamtram2/template_manage.html'
    form_class = TemplateForm

    def dispatch(self, request, *args, **kwargs):
        if not (
            request.user.groups.filter(name="WAMTRAM2_TEAM_LEADER").exists()
            or request.user.groups.filter(name="WAMTRAM2_STAFF").exists()
            or request.user.is_superuser
        ):
            return HttpResponseForbidden("You do not have permission to access this page.")
        
        elif request.method == 'DELETE':
            return self.delete(request, *args, **kwargs)
        elif request.method == 'GET' and 'location_code' in request.GET:
            return self.get_places(request)
        
        return super().dispatch(request, *args, **kwargs)

    def form_valid(self, form):
        form.save()
        messages.success(self.request, 'Template created successfully')
        
        if self.request.headers.get('x-requested-with') == 'XMLHttpRequest':
            return JsonResponse({
                'message': 'Template created successfully',
                'status': 'success'
            })
        return redirect('wamtram2:template_manage')

    def form_invalid(self, form):
        errors = []
        for field, error_list in form.errors.items():
            for error in error_list:
                errors.append(f'{field}: {error}')
        
        if self.request.headers.get('x-requested-with') == 'XMLHttpRequest':
            return JsonResponse({
                'message': 'Form submission failed',
                'errors': errors,
                'status': 'error'
            }, status=400)

        for error in errors:
            messages.error(self.request, error)
        return super().form_invalid(form)

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        context['templates'] = Template.objects.all().order_by('-template_id')
        context['locations'] = list(TrtLocations.get_ordered_locations())
        context['places'] = list(TrtPlaces.objects.all())
        context['species'] = list(TrtSpecies.objects.all())
        context['sex_choices'] = SEX_CHOICES
        return context

    def delete(self, request, template_key):
        template = get_object_or_404(Template, pk=template_key)
        template.delete()
        messages.success(request, 'Template deleted successfully')

        if request.headers.get('x-requested-with') == 'XMLHttpRequest':

            return JsonResponse({'message': 'Template deleted successfully', 'status': 'success'})
        
        return redirect('wamtram2:template_manage')
    
    def get_places(self, request):
        location_code = request.GET.get('location_code')
        places = TrtPlaces.objects.filter(location_code=location_code)
        places_list = list(places.values('place_code', 'place_name'))
        return JsonResponse(places_list, safe=False)


def get_place_full_name(request):
    place_code = request.GET.get('place_code')
    try:
        place = TrtPlaces.objects.get(place_code=place_code)
        full_name = place.get_full_name()
        return JsonResponse({'full_name': full_name})
    except TrtPlaces.DoesNotExist:
        return JsonResponse({'error': 'Place not found'}, status=404)
    
    
def check_template_name(request):
    name = request.GET.get('name')
    is_available = not Template.objects.filter(name=name).exists()
    return JsonResponse({'is_available': is_available})
    
    
class ValidateTagView(View):
    """
    View for validating tags.
    Provides functionality to validate different types of tags.

    Methods:
        validate_recaptured_tag(request): Validates a recaptured tag.
        validate_new_tag(request): Validates a new tag.
        validate_new_pit_tag(request): Validates a new PIT tag.
        validate_recaptured_pit_tag(request): Validates a recaptured PIT tag.
        get(request, *args, **kwargs): Handles GET requests.
    """

    def validate_recaptured_tag(self, request):
        """
        Validates a recaptured tag.

        Args:
            request (HttpRequest): The HTTP request.

        Returns:
            JsonResponse: The JSON response.
        """
        turtle_id = request.GET.get('turtle_id')
        tag = request.GET.get('tag', '')
        side = request.GET.get('side')

        if not tag or not side:
            return JsonResponse({'valid': False, 'wrong_side': False, 'message': 'Missing parameters'})
        
        if turtle_id:
            try:
                turtle_id = int(turtle_id)
                tag_obj = TrtTags.objects.filter(tag_id=tag).first()

                if tag_obj:
                    if tag_obj.turtle_id != turtle_id:
                        return JsonResponse({
                            'valid': False, 
                            'wrong_side': False, 
                            'message': 'Tag belongs to another turtle', 
                            'other_turtle_id': tag_obj.turtle_id,
                            'status': tag_obj.tag_status.description
                        })
                    else:
                        if tag_obj.tag_status.tag_status != 'ATT':
                            return JsonResponse({
                                'valid': False,
                                'wrong_side': False,
                                'message': f'Tag status: {tag_obj.tag_status.description}',
                                'status': tag_obj.tag_status.description
                            })
                        is_valid = True
                        wrong_side = (tag_obj.side.lower() != side.lower())
                        return JsonResponse({
                            'valid': is_valid, 
                            'wrong_side': wrong_side, 
                            'other_turtle_id': None,
                            'status': tag_obj.tag_status.description
                        })
            except TrtTurtles.DoesNotExist:
                    return JsonResponse({'valid': False, 'wrong_side': False, 'message': 'Turtle not found'})

        new_tag_entry = TrtDataEntry.objects.filter(
            Q(new_left_tag_id__tag_id=tag) |
            Q(new_left_tag_id_2__tag_id=tag) |
            Q(new_right_tag_id__tag_id=tag) |
            Q(new_right_tag_id_2__tag_id=tag),
            observation_id__isnull=True,
        ).order_by('-entry_batch__entry_date').first()
                
        if new_tag_entry:
            actual_side = None
            if new_tag_entry.new_left_tag_id and new_tag_entry.new_left_tag_id.tag_id == tag:
                actual_side = 'L'
            elif new_tag_entry.new_left_tag_id_2 and new_tag_entry.new_left_tag_id_2.tag_id == tag:
                actual_side = 'L'
            elif new_tag_entry.new_right_tag_id and new_tag_entry.new_right_tag_id.tag_id == tag:
                actual_side = 'R'
            elif new_tag_entry.new_right_tag_id_2 and new_tag_entry.new_right_tag_id_2.tag_id == tag:
                actual_side = 'R'
            
            if actual_side:
                wrong_side = (actual_side.lower() != side.lower())
            
            return JsonResponse({
                'valid': True, 
                'wrong_side': wrong_side,
                'message': 'Tag found in previous unprocessed entry',
                'entry_date': new_tag_entry.entry_batch.entry_date.strftime('%Y-%m-%d')
            })
        else:
            return JsonResponse({'valid': False, 'wrong_side': False, 'message': 'Tag not found', 'tag_not_found': True})

    def validate_new_tag(self, request):
        """
        Validates a new tag.

        Args:
            request (HttpRequest): The HTTP request.

        Returns:
            JsonResponse: The JSON response.
        """
        tag = request.GET.get('tag', '')
        if not tag:
            return JsonResponse({'valid': False, 'message': 'Missing tag parameter'})

        try:
            tag_obj = TrtTags.objects.filter(tag_id=tag).select_related('tag_status').first()
            if tag_obj:
                if tag_obj.turtle_id:
                    return JsonResponse({
                        'valid': False, 
                        'message': 'Tag belongs to another turtle', 
                        'other_turtle_id': tag_obj.turtle_id,
                        'status': tag_obj.tag_status.description
                    })
                if tag_obj.tag_status.tag_status == 'U':
                    return JsonResponse({'valid': True, 'other_turtle_id': None, 'status': tag_obj.tag_status.description})
                else:
                    return JsonResponse({'valid': False, 'status': tag_obj.tag_status.description})
            else:
                return JsonResponse({'valid': False, 'message': 'Flipper tag not found', 'tag_not_found': True})
        except Exception as e:
            return JsonResponse({'valid': False, 'message': str(e)})

    def validate_new_pit_tag(self, request):
        """
        Validates a new PIT tag.

        Args:
            request (HttpRequest): The HTTP request.

        Returns:
            JsonResponse: The JSON response.
        """
        tag = request.GET.get('tag')
        if not tag:
            return JsonResponse({'valid': False, 'message': 'Missing tag parameter'})

        try:
            pit_tag = TrtPitTags.objects.filter(pittag_id=tag).select_related('turtle', 'pit_tag_status').first()
            if pit_tag:
                if pit_tag.turtle:
                    return JsonResponse({
                        'valid': False, 
                        'message': 'PIT tag belongs to another turtle', 
                        'other_turtle_id': pit_tag.turtle.turtle_id,
                        'status': pit_tag.pit_tag_status.description
                    })
                if pit_tag.pit_tag_status.pit_tag_status == 'U':
                    return JsonResponse({'valid': True, 'other_turtle_id': None, 'status': pit_tag.pit_tag_status.description})
                else:
                    return JsonResponse({'valid': False, 'status': pit_tag.pit_tag_status.description})
            else:
                return JsonResponse({'valid': False, 'message': 'PIT tag not found', 'tag_not_found': True})
        except Exception as e:
            return JsonResponse({'valid': False, 'message': str(e)})

    def validate_recaptured_pit_tag(self, request):
        """
        Validates a recaptured PIT tag.
        """
        turtle_id = request.GET.get('turtle_id')
        tag = request.GET.get('tag')

        if not tag:
            return JsonResponse({'valid': False, 'message': 'Missing parameters'})

        try:
            new_pit_tag_entry = TrtDataEntry.objects.filter(
                Q(new_pittag_id__pittag_id=tag) |
                Q(new_pittag_id_2__pittag_id=tag) |
                Q(new_pittag_id_3__pittag_id=tag) |
                Q(new_pittag_id_4__pittag_id=tag),
                observation_id__isnull=True,
            ).order_by('-entry_batch__entry_date').first()
            
            if new_pit_tag_entry:
                return JsonResponse({
                    'valid': True, 
                    'message': 'PIT Tag found in previous unprocessed entry',
                    'entry_date': new_pit_tag_entry.entry_batch.entry_date.strftime('%Y-%m-%d')
                })

            if turtle_id:
                turtle_id = int(turtle_id)
                pit_tag = TrtPitTags.objects.filter(pittag_id=tag).select_related('turtle').first()

                if pit_tag:
                    if pit_tag.turtle and pit_tag.turtle.turtle_id != int(turtle_id):
                        return JsonResponse({
                            'valid': False,
                            'message': 'PIT tag belongs to another turtle',
                            'other_turtle_id': pit_tag.turtle.turtle_id,
                            'status': pit_tag.pit_tag_status.description
                        })
                    elif pit_tag.pit_tag_status.pit_tag_status != 'ATT':
                        return JsonResponse({
                            'valid': False,
                            'message': f'PIT tag status: {pit_tag.pit_tag_status.description}',
                            'status': pit_tag.pit_tag_status.description
                        })
                    else:
                        return JsonResponse({'valid': True})
                else:
                    return JsonResponse({'valid': False, 'message': 'PIT tag not found', 'tag_not_found': True})
            else:
                new_pit_tag_entry = TrtDataEntry.objects.filter(
                Q(new_pittag_id__pittag_id=tag) |
                Q(new_pittag_id_2__pittag_id=tag) |
                Q(new_pittag_id_3__pittag_id=tag) |
                Q(new_pittag_id_4__pittag_id=tag),
                observation_id__isnull=True,
                ).order_by('-entry_batch__entry_date').first()
                
                if new_pit_tag_entry:
                    return JsonResponse({
                        'valid': True, 
                        'message': 'Tag found in previous unprocessed entry',
                        'entry_date': new_pit_tag_entry.entry_batch.entry_date.strftime('%Y-%m-%d')
                    })
                else:
                    return JsonResponse({'valid': False, 'message': 'PIT tag not found', 'tag_not_found': True})
        except Exception as e:
            return JsonResponse({'valid': False, 'message': str(e)})

    def get(self, request, *args, **kwargs):
        """
        Handles GET requests for tag validation.

        Args:
            request (HttpRequest): The HTTP request.

        Returns:
            JsonResponse: The JSON response.
        """
        validation_type = request.GET.get('type')
        if validation_type == 'recaptured_tag':
            return self.validate_recaptured_tag(request)
        elif validation_type == 'new_tag':
            return self.validate_new_tag(request)
        elif validation_type == 'new_pit_tag':
            return self.validate_new_pit_tag(request)
        elif validation_type == 'recaptured_pit_tag':
            return self.validate_recaptured_pit_tag(request)
        else:
            return JsonResponse({'valid': False, 'message': 'Invalid validation type'})


def search_persons(request):
    query = request.GET.get('q', '')
    query_parts = query.split()

    # Build the query
    if len(query_parts) == 1:
        persons = TrtPersons.objects.filter(
            Q(first_name__icontains=query_parts[0]) | Q(surname__icontains=query_parts[0])
        ).values('person_id', 'first_name', 'surname')
    elif len(query_parts) >= 2:
        persons = TrtPersons.objects.filter(
            Q(first_name__icontains=query_parts[0]) & Q(surname__icontains=query_parts[1])
        ).values('person_id', 'first_name', 'surname')

    return JsonResponse(list(persons), safe=False)


def search_places(request):
    query = request.GET.get('q', '')
    
    if '(' in query and ')' in query:
        place_name, location_code = query.split('(', 1)
        location_code = location_code.rstrip(')')
        places = TrtPlaces.objects.filter(
            place_name__icontains=place_name.strip(),
            location_code__location_name__icontains=location_code.strip()
        ).values('place_code', 'place_name', 'location_code__location_name')[:10]
    else:
        places = TrtPlaces.objects.filter(
            Q(place_name__icontains=query) | Q(location_code__location_name__icontains=query)
        ).values('place_code', 'place_name', 'location_code__location_name')[:10]
        
    for place in places:
        place['full_name'] = f"{place['place_name']} ({place['location_code__location_name']})"
    
    return JsonResponse(list(places), safe=False)


class ExportDataView(LoginRequiredMixin, View):
    template_name = 'wamtram2/export_form.html'

    def _get_date_range(self, date_from, date_to):
        """
        Convert date strings to timezone-aware datetime objects
        Args:
            date_from: YYYY-MM-DD string
            date_to: YYYY-MM-DD string
        Returns:
            tuple of (start_datetime, end_datetime)
        """
        if not date_from or not date_to:
            return None, None
            
        start_date = timezone.make_aware(
            datetime.combine(
                datetime.strptime(date_from, '%Y-%m-%d').date(),
                time.min
            )
        )
        end_date = timezone.make_aware(
            datetime.combine(
                datetime.strptime(date_to, '%Y-%m-%d').date(),
                time.max
            )
        )
        return start_date, end_date

    def dispatch(self, request, *args, **kwargs):
        # Permission check: only allow users in the specific groups or superusers
        if not (
            request.user.groups.filter(name="WAMTRAM2_STAFF").exists()
            or request.user.groups.filter(name="WAMTRAM2_TEAM_LEADER").exists()
            or request.user.is_superuser
        ):
            return HttpResponseForbidden(
                "You do not have permission to view this record"
            )
        return super().dispatch(request, *args, **kwargs)

    def get(self, request):
        
        # If it's an export request (has format parameter)
        if request.GET.get("format"):
            return self.export_data(request)
            
        # Handle different actions based on the 'action' parameter
        action = request.GET.get('action')
        if action:
            if action == 'get_places':
                return self.get_places(request)
            elif action == 'get_species':
                return self.get_species(request)
            elif action == 'get_sexes':
                return self.get_sexes(request)
        
        # If no action or format, render the form
        return render(request, self.template_name)

    def get_places(self, request):
        """Retrieve places based on the specified date range."""
        from_date, to_date = self._get_date_range(
            request.GET.get("observation_date_from"),
            request.GET.get("observation_date_to")
        )

        # First get data accessible to the user
        queryset = TrtDataEntry.objects.all()
        user = request.user
        if not user.is_superuser:
            user_organisations = user.organisations.all()
            if user_organisations.exists():
                related_batch_ids = TrtEntryBatchOrganisation.objects.filter(
                    organisation__in=[org.code for org in user_organisations]
                ).values_list('trtentrybatch_id', flat=True)
                queryset = queryset.filter(entry_batch_id__in=related_batch_ids)
            else:
                return JsonResponse({"places": []})

        if from_date and to_date:
            queryset = queryset.filter(observation_date__range=[from_date, to_date])

        places = TrtPlaces.objects.filter(
            place_code__in=queryset.values_list('place_code', flat=True)
        ).select_related('location_code').distinct()

        place_list = [
            {
                "value": place.place_code,
                "label": place.get_full_name(),
                "location_name": place.location_code.location_name,
            }
            for place in places
        ]
        
        return JsonResponse({"places": place_list})

    def get_species(self, request):
        """Retrieve species based on the specified date range."""
        from_date, to_date = self._get_date_range(
            request.GET.get("observation_date_from"),
            request.GET.get("observation_date_to")
        )

        queryset = TrtDataEntry.objects.all()
        user = request.user
        if not user.is_superuser:
            user_organisations = user.organisations.all()
            if user_organisations.exists():
                related_batch_ids = TrtEntryBatchOrganisation.objects.filter(
                    organisation__in=[org.code for org in user_organisations]
                ).values_list('trtentrybatch_id', flat=True)
                queryset = queryset.filter(entry_batch_id__in=related_batch_ids)
            else:
                return JsonResponse({"species": []})

        if from_date and to_date:
            queryset = queryset.filter(observation_date__range=[from_date, to_date])

        species = TrtSpecies.objects.filter(
            species_code__in=queryset.values_list('species_code', flat=True)
        ).distinct()

        species_list = [
            {"value": specie.species_code, "label": specie.common_name}
            for specie in species
        ]

        return JsonResponse({"species": species_list})

    def get_sexes(self, request):
        """Retrieve available sex choices based on the defined SEX_CHOICES."""
        from_date, to_date = self._get_date_range(
            request.GET.get("observation_date_from"),
            request.GET.get("observation_date_to")
        )

        queryset = TrtDataEntry.objects.all()
        user = request.user
        if not user.is_superuser:
            user_organisations = user.organisations.all()
            if user_organisations.exists():
                related_batch_ids = TrtEntryBatchOrganisation.objects.filter(
                    organisation__in=[org.code for org in user_organisations]
                ).values_list('trtentrybatch_id', flat=True)
                queryset = queryset.filter(entry_batch_id__in=related_batch_ids)
            else:
                return JsonResponse({"sexes": []})

        if from_date and to_date:
            queryset = queryset.filter(observation_date__range=[from_date, to_date])

        used_sexes = queryset.values_list('sex', flat=True).distinct()
        sex_list = [
            {"value": choice[0], "label": choice[1]}
            for choice in SEX_CHOICES
            if choice[0] in used_sexes
        ]

        return JsonResponse({"sexes": sex_list})

    def export_data(self, request):
        try:
            from_date, to_date = self._get_date_range(
                request.GET.get("observation_date_from"),
                request.GET.get("observation_date_to")
            )
            
            place_code = request.GET.get("place_code")
            species = request.GET.get("species")
            sex = request.GET.get("sex")
            file_format = request.GET.get("format", "csv")
            
            filename_parts = []
            if place_code:
                filename_parts.append(place_code)
            if species:
                filename_parts.append(species)
            if sex:
                filename_parts.append(sex)
                
            date_range = ""
            if from_date and to_date:
                date_range = f"({from_date.strftime('%Y%m%d')}-{to_date.strftime('%Y%m%d')})"
            
            filename = "_".join(filename_parts) + date_range
            
            if not filename:
                filename = f"data_export{date_range}"

            queryset = TrtDataEntry.objects.all()
            
            # Apply organization filter
            user = request.user
            if not user.is_superuser:
                user_organisations = user.organisations.all()
                if user_organisations.exists():
                    related_batch_ids = TrtEntryBatchOrganisation.objects.filter(
                        organisation__in=[org.code for org in user_organisations]
                    ).values_list('trtentrybatch_id', flat=True)
                    queryset = queryset.filter(entry_batch_id__in=related_batch_ids)
                else:
                    return HttpResponse("No data available for your organisation")
            
            # Apply filters
            if from_date and to_date:
                queryset = queryset.filter(observation_date__range=[from_date, to_date])
            if place_code:
                queryset = queryset.filter(place_code=place_code)
            if species:
                queryset = queryset.filter(species_code=species)
            if sex:
                queryset = queryset.filter(sex=sex)
                
            queryset = queryset.select_related('entry_batch')

            if file_format == "csv":
                response = HttpResponse(content_type="text/csv")
                response["Content-Disposition"] = f'attachment; filename="{filename}.csv"'  # 注意这里的f-string
                writer = csv.writer(response)

                headers = [field.name for field in TrtDataEntry._meta.fields]
                headers.append('organisations')
                writer.writerow(headers)
        
                
                for entry in queryset:
                    organisations = TrtEntryBatchOrganisation.objects.filter(
                        trtentrybatch=entry.entry_batch
                    ).values_list('organisation', flat=True)
                    org_str = ', '.join(organisations)
                    
                    row = [getattr(entry, field.name) for field in TrtDataEntry._meta.fields]
                    row.append(org_str)
                    writer.writerow(row)
                    
            else:  # xlsx format
                response = HttpResponse(content_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")
                response["Content-Disposition"] = f'attachment; filename="{filename}.xlsx"'  # 注意这里的f-string
                wb = Workbook()
                ws = wb.active
                
                headers = [field.name for field in TrtDataEntry._meta.fields]
                headers.append('organisations')
                ws.append(headers)
                
                for entry in queryset:
                    organisations = TrtEntryBatchOrganisation.objects.filter(
                        trtentrybatch=entry.entry_batch
                    ).values_list('organisation', flat=True)
                    org_str = ', '.join(organisations)
                    
                    row = []
                    for field in TrtDataEntry._meta.fields:
                        value = getattr(entry, field.name)
                        if field.is_relation:
                            value = str(value) if value else ''
                        elif isinstance(value, datetime):
                            value = value.strftime('%Y-%m-%d %H:%M:%S')
                        elif isinstance(value, date):
                            value = value.strftime('%Y-%m-%d')
                        elif value is None:
                            value = ''
                        row.append(value)
                    
                    row.append(org_str)
                    ws.append(row)
                
                wb.save(response)

            return response

        except Exception as e:
            import traceback
            traceback.print_exc()
            return HttpResponse(f"Error during export: {str(e)}", status=500)


class DudTagManageView(LoginRequiredMixin, View):
    template_name = 'wamtram2/dud_tag_manage.html'

    def dispatch(self, request, *args, **kwargs):
        if not request.user.is_superuser:
            return HttpResponseForbidden("You do not have permission to view this record")
        return super().dispatch(request, *args, **kwargs)

    def get(self, request):
        entries = TrtDataEntry.objects.filter(
            dud_flipper_tag__isnull=False
        ).select_related('observation_id', 'turtle_id')
        tag_states = TrtTagStates.objects.all()
        return render(request, self.template_name, {'entries': entries, 'tag_states': tag_states})

    def post(self, request):
        entry_id = request.POST.get('entry_id')
        tag_type = request.POST.get('tag_type')
        tag_id = request.POST.get('tag_id')
        entry = get_object_or_404(TrtDataEntry, pk=entry_id)

        if entry.observation_id:
            observation = entry.observation_id
            
            if tag_type == 'flipper':
                observation.dud_flipper_tag = tag_id
            elif tag_type == 'flipper_2':
                observation.dud_flipper_tag_2 = tag_id
            elif tag_type == 'pit':
                observation.dud_pit_tag = tag_id
            elif tag_type == 'pit_2':
                observation.dud_pit_tag_2 = tag_id

            observation.save()

        return redirect('wamtram2:dud_tag_manage')


class BatchesCurationView(LoginRequiredMixin,ListView):
    model = TrtEntryBatches
    template_name = 'wamtram2/batches_curation.html'
    context_object_name = 'batches'
    paginate_by = 20
    
    def dispatch(self, request, *args, **kwargs):
        if not (
            request.user.groups.filter(name="WAMTRAM2_TEAM_LEADER").exists()
            or request.user.groups.filter(name="WAMTRAM2_STAFF").exists()
            or request.user.is_superuser
        ):
            return HttpResponseForbidden("You do not have permission to view this record")
        return super().dispatch(request, *args, **kwargs)

    def get_queryset(self):
        user_organisations = self.request.user.organisations.all()
        queryset = super().get_queryset().order_by('-entry_batch_id')
        user = self.request.user
        
        if user.is_superuser:
            pass
        elif not user_organisations.exists():
            return queryset.none()
        else:
            for org in user_organisations:
                related_batch_ids = TrtEntryBatchOrganisation.objects.filter(
                    organisation=org.code
                ).values_list('trtentrybatch_id', flat=True)

            queryset = TrtEntryBatches.objects.filter(
                entry_batch_id__in=related_batch_ids
            ).order_by('-entry_batch_id')

        last_place_code_subquery = Subquery(
            TrtDataEntry.objects.filter(entry_batch_id=OuterRef("pk"))
            .order_by("-data_entry_id")
            .values("place_code__place_name")[:1]
        )

        queryset = queryset.annotate(
            entry_count=Count('trtdataentry'),
            last_place_code=last_place_code_subquery,
            do_not_process_count=Count(
                'trtdataentry',
                filter=Q(trtdataentry__do_not_process=True)
            )
        )
        
        location = self.request.GET.get('location')
        place = self.request.GET.get('place')
        year = self.request.GET.get('year')
        show_all = self.request.GET.get('show_all')
        
        if not self.request.GET:
            return queryset[:20]
        
        if show_all:
            return queryset
        
        if not (location or year):
            return TrtEntryBatches.objects.none()

        query = Q()

        if location and place and year:
            year_code = str(year)[-2:]
            query = Q(batches_code__contains=place) & Q(batches_code__endswith=year_code)
        elif location and year:
            year_code = str(year)[-2:]
            query = Q(batches_code__contains=location) & Q(batches_code__endswith=year_code)
        elif location:
            query = Q(batches_code__contains=location)
        elif year:
            year_code = str(year)[-2:]
            query = Q(batches_code__endswith=year_code)

        result = queryset.filter(query).order_by('-entry_batch_id') if query else queryset.order_by('-entry_batch_id')
        return result
    
    
    def get_user_role(self, user):
        if user.is_superuser:
            return "Super User"
        user_groups = user.groups.values_list('name', flat=True)
        if "WAMTRAM2_STAFF" in user_groups:
            return "Staff"
        elif "WAMTRAM2_TEAM_LEADER" in user_groups:
            return "Team Leader"
        elif "WAMTRAM2_VOLUNTEER" in user_groups:
            return "Volunteer"
        
        return ""

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        
        user = self.request.user
        
        role = self.get_user_role(user)
        context['user_role'] = role
        
        user_organisations = self.request.user.organisations.all()
        context['user_organisation_codes'] = [org.code for org in user_organisations]
        
        context['locations'] = list(TrtLocations.get_ordered_locations())
        context['places'] = TrtPlaces.objects.all().order_by('place_name')
        current_year = timezone.now().year
        context['years'] = range(2022, current_year + 1)
        context['selected_location'] = self.request.GET.get('location', '')
        context['selected_place'] = self.request.GET.get('place', '')
        context['selected_year'] = self.request.GET.get('year', '')
        context['templates'] = Template.objects.all()
        context['batches'] = self.get_queryset()
        context['is_initial_load'] = not bool(self.request.GET)
        places = TrtPlaces.objects.select_related('location_code').all()
        places_data = [
            {
                'place_code': place.place_code,
                'place_name': place.place_name,
                'full_name': place.get_full_name()
            }
            for place in places
        ]
        context['places_json'] = json.dumps(places_data, cls=DjangoJSONEncoder)
        context['show_all'] = self.request.GET.get('show_all', False)

        return context
    def get(self, request, *args, **kwargs):
        if 'action' in request.GET:
            action = request.GET.get('action')
            if action == 'check_batch_code':
                return self.check_batch_code(request)
            
        self.object_list = self.get_queryset()
        context = self.get_context_data()
        
        if request.headers.get('X-Requested-With') == 'XMLHttpRequest':
            html = render_to_string('wamtram2/batches_curation.html', context, request=request)
            return JsonResponse({
                'html': html,
                'count': context['paginator'].count if 'paginator' in context else 0,
                'num_pages': context['paginator'].num_pages if 'paginator' in context else 1,
                'current_page': context['page_obj'].number if 'page_obj' in context else 1,
                'show_all': context.get('show_all', False),
            })
        return super().get(request, *args, **kwargs)
    
    def check_batch_code(self, request):
        code = request.GET.get('code')
        batch_id = request.GET.get('batch_id')
        if batch_id:
            is_unique = not TrtEntryBatches.objects.filter(batches_code=code).exclude(pk=batch_id).exists()
        else:
            is_unique = not TrtEntryBatches.objects.filter(batches_code=code).exists()
        return JsonResponse({'is_unique': is_unique})
    
    def get_places(self, request):
        location_code = request.GET.get('location_code')
        places = TrtPlaces.objects.filter(location_code=location_code).values('place_code', 'place_name')
        return JsonResponse(list(places), safe=False)
    
    
class CreateNewEntryView(LoginRequiredMixin, ListView):
    model = TrtEntryBatches
    template_name = 'wamtram2/create_new_entry.html'
    context_object_name = 'batches'
    paginate_by = 20
    
    def dispatch(self, request, *args, **kwargs):
        # Permission check: only allow users in the specific groups or superusers
        if not (
            request.user.groups.filter(name="WAMTRAM2_VOLUNTEER").exists()
            or request.user.groups.filter(name="WAMTRAM2_TEAM_LEADER").exists()
            or request.user.groups.filter(name="WAMTRAM2_STAFF").exists()
            or request.user.is_superuser
        ):
            return HttpResponseForbidden(
                "You do not have permission to view this record"
            )
        return super().dispatch(request, *args, **kwargs)
    
    def get_user_role(self, user):
        if user.is_superuser:
            return "Super User"
        user_groups = user.groups.values_list('name', flat=True)
        if "WAMTRAM2_STAFF" in user_groups:
            return "Staff"
        elif "WAMTRAM2_TEAM_LEADER" in user_groups:
            return "Team Leader"
        elif "WAMTRAM2_VOLUNTEER" in user_groups:
            return "Volunteer"
        
        return ""

    def get_queryset(self):
        """
        Filter the batches data based on query parameters
        """
        queryset = super().get_queryset().order_by('-entry_batch_id')

        user = self.request.user
        if not user.is_superuser:
            user_organisations = self.request.user.organisations.all()
            if not user_organisations.exists():
                return queryset.none()

            related_batch_ids = TrtEntryBatchOrganisation.objects.filter(
                organisation__in=[org.code for org in user_organisations]
            ).values_list('trtentrybatch_id', flat=True)

            queryset = queryset.filter(entry_batch_id__in=related_batch_ids)
        
        if self.request.GET.get('show_all'):
            return queryset

        location = self.request.GET.get('location')
        place = self.request.GET.get('place')
        year = self.request.GET.get('year')

        query = Q()
        current_year = timezone.now().year
        years = {str(year): str(year)[-2:] for year in range(2022, current_year + 1)}

        # Generate queries based on different filter parameters
        if location and place and year:
            year_code = years.get(year)
            if year_code:
                query = Q(batches_code__contains=place) & Q(batches_code__endswith=year_code)
        elif location and year:
            year_code = years.get(year)
            if year_code:
                query = Q(batches_code__contains=location) & Q(batches_code__endswith=year_code)
        elif location:
            query = Q(batches_code__contains=location)
        elif year:
            year_code = years.get(year)
            if year_code:
                query = Q(batches_code__endswith=year_code)

        return queryset.filter(query).order_by('-entry_batch_id')


    def get_context_data(self, **kwargs):
        """
        Provide context data to the template, including locations, places, and years
        """
        context = super().get_context_data(**kwargs)
        
        user = self.request.user
        
        role = self.get_user_role(user)
        context['user_role'] = role

        user_organisations = self.request.user.organisations.all()
        context['user_organisation_codes'] = [org.code for org in user_organisations]

        locations = TrtLocations.get_ordered_locations()
        places = TrtPlaces.objects.none()

        if 'location' in self.request.GET and self.request.GET['location']:
            places = TrtPlaces.objects.filter(location_code=self.request.GET['location'])

        current_year = timezone.now().year
        years = {str(year): str(year)[-2:] for year in range(2022, current_year + 1)}

        context.update({
            'locations': locations,
            'places': places,
            'years': years,
            'selected_location': self.request.GET.get('location', ''),
            'selected_place': self.request.GET.get('place', ''),
            'selected_year': self.request.GET.get('year', ''),
            'templates': Template.objects.all(),
        })

        return context

    def get(self, request, *args, **kwargs):
        self.object_list = self.get_queryset()
        context = self.get_context_data()
        
        if request.headers.get('X-Requested-With') == 'XMLHttpRequest':
            html = render_to_string('wamtram2/create_new_entry.html', context, request=request)
            return JsonResponse({
                'html': html,
                'count': context['paginator'].count if 'paginator' in context else 0,
                'num_pages': context['paginator'].num_pages if 'paginator' in context else 1,
                'current_page': context['page_obj'].number if 'page_obj' in context else 1,
            })
        return super().get(request, *args, **kwargs)
    

@login_required
@require_POST
def quick_add_batch(request):
    batches_code = request.POST.get('batches_code')
    comments = request.POST.get('comments', '')
    template_id = request.POST.get('template')
    entered_person_id = request.POST.get('entered_person_id')

    entered_person = None
    if entered_person_id:
        try:
            entered_person = TrtPersons.objects.get(pk=entered_person_id)
        except TrtPersons.DoesNotExist:
            return JsonResponse({'success': False, 'error': 'Invalid entered person ID.'})
    
    template = None
    if template_id:
        try:
            template = Template.objects.get(pk=template_id)
        except Template.DoesNotExist:
            return JsonResponse({'success': False, 'error': 'Invalid template ID.'})
    
    try:
        batch = TrtEntryBatches.objects.create(
            batches_code=batches_code,
            comments=comments,
            entry_date=timezone.now(),
            pr_date_convention=False,
            entered_person_id=entered_person,
            template=template
        )
        
        user_organisations = request.user.organisations.all()
        
        for org in user_organisations:
            TrtEntryBatchOrganisation.objects.create(
                trtentrybatch=batch,
                organisation=org.code
            )
        
        return JsonResponse({'success': True})
    except ValidationError as e:
        return JsonResponse({'success': False, 'error': str(e)})
    except Exception as e:
        return JsonResponse({'success': False, 'error': str(e)})


def search_templates(request):
    query = request.GET.get('q', '')
    if len(query) >= 2:
        templates = Template.objects.filter(
            Q(name__icontains=query)
        )[:10]
        data = [{'template_id': t.template_id, 'name': t.name} for t in templates]
        return JsonResponse(data, safe=False)
    return JsonResponse([], safe=False)
    
    
class BatchCodeManageView(View):
    template_name = 'wamtram2/batch_detail_manage.html'

    def dispatch(self, request, *args, **kwargs):
        
        if not (
            request.user.groups.filter(name="WAMTRAM2_TEAM_LEADER").exists()
            or request.user.groups.filter(name="WAMTRAM2_STAFF").exists()
            or request.user.is_superuser
        ):
            return HttpResponseForbidden(
                "You do not have permission to view this record"
            )
            
        if request.method == 'GET' and 'action' in request.GET:
            action = request.GET.get('action')
            if action == 'get_places':
                return self.get_places(request)
            elif action == 'check_batch_code':
                return self.check_batch_code(request)
        return super().dispatch(request, *args, **kwargs)
    
    def get_places(self, request):
        location_code = request.GET.get('location_code')
        places = TrtPlaces.objects.filter(location_code=location_code).values('place_code', 'place_name')
        return JsonResponse(list(places), safe=False)

    def get(self, request, batch_id=None):
        if batch_id:
            batch = get_object_or_404(TrtEntryBatches, pk=batch_id)
            form = BatchesCodeForm(instance=batch)
            entered_person = batch.entered_person_id
            entered_person_full_name = str(entered_person) if entered_person else ''
            entered_person_id = entered_person.person_id if entered_person else ''
            template = batch.template
            template_name = template.name if template else ''
            template_id = template.template_id if template else ''
        else:
            form = BatchesCodeForm()
            entered_person_full_name = ''
            entered_person_id = ''

        locations = TrtLocations.get_ordered_locations()
        current_year = timezone.now().year
        years = {str(year): str(year)[-2:] for year in range(2022, current_year + 1)}
        templates = Template.objects.all()

        entered_person_full_name = str(form.instance.entered_person_id) if form.instance.entered_person_id else ''
        context = {
            'form': form,
            'locations': locations,
            'years': years,
            'current_year': current_year,
            'templates': templates,
            'batch_id': batch_id,
            'entered_person_full_name': entered_person_full_name,
            'entered_person_id': entered_person_id,
            'template_name': template_name,
            'template_id': template_id,
        }
        return render(request, self.template_name, context)

    def post(self, request, batch_id=None):
        if batch_id:
            batch = get_object_or_404(TrtEntryBatches, pk=batch_id)
            form = BatchesCodeForm(request.POST, instance=batch)
        else:
            form = BatchesCodeForm(request.POST)

        if form.is_valid():
            new_batch = form.save(commit=False)
            if not TrtEntryBatches.objects.filter(batches_code=new_batch.batches_code).exclude(pk=batch_id).exists():
                new_batch.entered_person_id = form.cleaned_data['entered_person_id']
                new_batch.template = form.cleaned_data['template']
                new_batch.save()
                return redirect(reverse('wamtram2:batches_curation'))
            else:
                form.add_error('batches_code', 'This batch code already exists.')

        locations = TrtLocations.get_ordered_locations()
        current_year = timezone.now().year
        years = {str(year): str(year)[-2:] for year in range(2022, current_year + 1)}
        templates = Template.objects.all()

        context = {
            'form': form,
            'locations': locations,
            'years': years,
            'current_year': current_year,
            'templates': templates,
            'batch_id': batch_id,
        }
        return render(request, self.template_name, context)

    @method_decorator(require_http_methods(["GET"]))
    def check_batch_code(self, request):
        code = request.GET.get('code')
        batch_id = request.GET.get('batch_id')
        if batch_id:
            is_unique = not TrtEntryBatches.objects.filter(batches_code=code).exclude(pk=batch_id).exists()
        else:
            is_unique = not TrtEntryBatches.objects.filter(batches_code=code).exists()
        return JsonResponse({'is_unique': is_unique})

    
@require_GET
def get_places(request):
    location_code = request.GET.get('location_code')
    places = TrtPlaces.objects.filter(location_code=location_code).values('place_code', 'place_name')
    return JsonResponse(list(places), safe=False)


class AddPersonView(LoginRequiredMixin, FormView):
    template_name = 'wamtram2/add_person.html'
    form_class = TrtPersonsForm

    def dispatch(self, request, *args, **kwargs):
        if not (
            request.user.groups.filter(name="WAMTRAM2_TEAM_LEADER").exists()
            or request.user.groups.filter(name="WAMTRAM2_STAFF").exists()
            or request.user.is_superuser
        ):
            raise PermissionDenied("You do not have permission to access this page.")
        return super().dispatch(request, *args, **kwargs)
    
    def get_form(self, form_class=None):
        form = super().get_form(form_class)
        field_order = [
            'first_name', 'surname', 'email', 'recorder',
            'middle_name', 'specialty', 'address_line_1', 'address_line_2',
            'town', 'state', 'post_code', 'country',
            'telephone', 'fax', 'mobile', 'comments', 'transfer'
        ]
        form.order_fields(field_order)
        return form

    def form_valid(self, form):
        form.cleaned_data['recorder'] = form.cleaned_data.get('recorder', False)
        form.save()
        messages.success(self.request, 'Person added!')

        if self.request.headers.get('x-requested-with') == 'XMLHttpRequest':
            return JsonResponse({
                'message': 'Person added!',
                'status': 'success'
            })
        return redirect('wamtram2:add_person')

    def form_invalid(self, form):
        errors = []
        for field, error_list in form.errors.items():
            for error in error_list:
                errors.append(f'{field}: {error}')
        
        if self.request.headers.get('x-requested-with') == 'XMLHttpRequest':
            return JsonResponse({
                'message': 'Form submission failed',
                'errors': errors,
                'status': 'error'
            }, status=400)

        for error in errors:
            messages.error(self.request, error)
        return super().form_invalid(form)

    def post(self, request, *args, **kwargs):
        if 'file' in request.FILES:
            return self.handle_file_upload(request)
        return super().post(request, *args, **kwargs)
    
    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        return context
    
    def handle_file_upload(self, request):
        file = request.FILES.get('file')
        if file:
            try:
                df = pd.read_excel(file) if file.name.endswith(('.xls', '.xlsx')) else pd.read_csv(file)
                existing_emails = set(TrtPersons.objects.values_list('email', flat=True))
                new_entries = []
                skipped_emails = []

                for _, row in df.iterrows():
                    email = row['email']
                    if email in existing_emails:
                        skipped_emails.append(email)
                        continue

                    new_entries.append(TrtPersons(
                        first_name=row['first_name'],
                        surname=row['surname'],
                        email=email,
                        recorder=row.get('recorder', False)
                    ))

                TrtPersons.objects.bulk_create(new_entries)
                added_count = len(new_entries)
                skipped_count = len(skipped_emails)
                
                messages.success(request, f'{added_count} people added!')

                if skipped_count > 0:
                    skipped_emails_str = ', '.join(skipped_emails)
                    messages.warning(request, f'The following emails already exist and were skipped: {skipped_emails_str}')

                if request.headers.get('x-requested-with') == 'XMLHttpRequest':
                    response_data = {
                        'message': f'{added_count} people added!',
                        'status': 'success'
                    }
                    if skipped_count > 0:
                        response_data['skipped_emails'] = skipped_emails
                        response_data['warning'] = f'{skipped_count} email(s) already existed and were skipped.'
                    
                    return JsonResponse(response_data)

            except Exception as e:
                error_message = f'Error: {str(e)}'
                messages.error(request, error_message)
                if request.headers.get('x-requested-with') == 'XMLHttpRequest':
                    return JsonResponse({
                        'message': error_message,
                        'status': 'error'
                    }, status=400)
        else:
            error_message = 'Please select a file'
            messages.error(request, error_message)
            if request.headers.get('x-requested-with') == 'XMLHttpRequest':
                return JsonResponse({
                    'message': error_message,
                    'status': 'error'
                }, status=400)

        return redirect('wamtram2:add_person')

class AvailableBatchesView(LoginRequiredMixin, View):
    def get(self, request):
        if request.user.is_superuser:
            batches = TrtEntryBatches.objects.all()
        else:
            user_orgs = request.user.organisations.all()
            batches = TrtEntryBatches.objects.filter(
                batch_organisations__organisation__in=[org.code for org in user_orgs]
            )
            
        current_batch_id = request.GET.get('current_batch_id')
        
        batches = batches.exclude(
            entry_batch_id=current_batch_id
        ).distinct()
        
        return JsonResponse([{
            'id': batch.entry_batch_id,
            'code': batch.batches_code,
            'comment': batch.comments
        } for batch in batches], safe=False)

class BatchInfoView(LoginRequiredMixin, View):
    def get(self, request, batch_id):
        try:
            user_org_codes = [org.code for org in request.user.organisations.all()]
            batch = TrtEntryBatches.objects.get(
                entry_batch_id=batch_id,
                batch_organisations__organisation__in=user_org_codes
            )
            return JsonResponse({
                'code': batch.batches_code,
                'comment': batch.comments
            })
        except TrtEntryBatches.DoesNotExist:
            return JsonResponse({'error': 'Batch not found'}, status=404)

class MoveEntryView(LoginRequiredMixin, View):
    def post(self, request):
        entry_id = request.POST.get('entry_id')
        target_batch_id = request.POST.get('target_batch_id')
        
        if not entry_id or not target_batch_id:
            return JsonResponse({'error': 'Missing required parameters'}, status=400)
            
        try:
            user_org_codes = [org.code for org in request.user.organisations.all()]
            
            entry = TrtDataEntry.objects.select_related('entry_batch').get(
                data_entry_id=entry_id
            )
            current_batch = entry.entry_batch
            
            if not current_batch.batch_organisations.filter(
                organisation__in=user_org_codes
            ).exists():
                raise PermissionDenied('No permission to operate on this entry')
            
            target_batch = TrtEntryBatches.objects.get(
                entry_batch_id=target_batch_id,
                batch_organisations__organisation__in=user_org_codes
            )
            
            entry.entry_batch = target_batch
            entry.save()
            
            return JsonResponse({
                'success': True,
                'message': f'Successfully moved entry to batch {target_batch.batches_code}'
            })
            
        except TrtDataEntry.DoesNotExist:
            return JsonResponse({'error': 'Entry not found'}, status=404)
        except TrtEntryBatches.DoesNotExist:
            return JsonResponse({'error': 'Target batch not found'}, status=404)
        except PermissionDenied as e:
            return JsonResponse({'error': str(e)}, status=403)
        except Exception as e:
            return JsonResponse({'error': f'Operation failed: {str(e)}'}, status=500)
        

class PersonManageView(LoginRequiredMixin, UserPassesTestMixin, PaginateMixin, ListView):
    model = TrtPersons
    template_name = 'wamtram2/manage_person.html'
    context_object_name = 'persons'
    paginate_by = 50
    
    def test_func(self):
        return self.request.user.is_superuser
    
    def handle_no_permission(self):
        if not self.request.user.is_authenticated:
            return super().handle_no_permission()
        raise PermissionDenied("You must be a superuser to access this page.")

    def get_queryset(self):
        queryset = super().get_queryset()
        search_term = self.request.GET.get('search', '').strip()
        
        if search_term:
            search_terms = search_term.split()
            query = Q()
            
            if len(search_terms) == 1:
                term = search_terms[0]
                query = (
                    Q(first_name__icontains=term) |
                    Q(surname__icontains=term) |
                    Q(email__icontains=term)
                )

            else:
                query = (
                    Q(first_name__icontains=search_term) |
                    Q(surname__icontains=search_term) |
                    Q(email__icontains=search_term) |
                    
                    Q(first_name__icontains=search_terms[0], 
                    surname__icontains=' '.join(search_terms[1:])) |
                    
                    reduce(operator.or_, (
                        Q(first_name__icontains=term) |
                        Q(surname__icontains=term) |
                        Q(email__icontains=term)
                        for term in search_terms
                    ))
                )
            
            queryset = queryset.filter(query)
    
        return queryset.prefetch_related(
            'measurer_person',
            'tagger_person',
            'entered_by_person'
        )

    @transaction.atomic
    def post(self, request, *args, **kwargs):
        action = request.POST.get('action')
        
        if action == 'merge':
            return self.handle_merge()
        elif action == 'update':
            return self.handle_update()
        
        return self.get(request, *args, **kwargs)

    def handle_merge(self):
        primary_id = self.request.POST.get('primary_person')
        secondary_id = self.request.POST.get('secondary_person')
        
        try:
            primary = TrtPersons.objects.get(pk=primary_id)
            secondary = TrtPersons.objects.get(pk=secondary_id)
            
            if primary == secondary:
                messages.error(self.request, "Cannot merge a person with themselves")
                return self.get(self.request)
                
            models_to_update = [
                ('TrtDataEntry', ['measured_by_id', 'recorded_by_id', 'tagged_by_id', 
                                'entered_by_id', 'measured_recorded_by_id']),
                ('TrtObservations', ['reporter_person', 'entered_by_person','tagger_person',
                                'measurer_reporter_person','measurer_person']),
            ]
            
            for model_name, fields in models_to_update:
                model = apps.get_model('wamtram2', model_name)
                for field in fields:
                    model.objects.filter(**{field: secondary}).update(**{field: primary})
            
            secondary_info = f"{secondary.first_name} {secondary.surname}"
            merge_note = f"Merged with {secondary_info} on {timezone.now().strftime('%Y-%m-%d')}"
            
            if primary.comments:
                primary.comments += f"\n{merge_note}"
            else:
                primary.comments = merge_note
            primary.save()
            
            secondary.delete()
            messages.success(self.request, 
                        f"Successfully merged {secondary_info} into {primary.first_name} {primary.surname}")
            
        except Exception as e:
            messages.error(self.request, f"Error during merge: {str(e)}")
        
        return self.get(self.request)

    def handle_update(self):
        person_id = self.request.POST.get('person_id')
        try:
            person = TrtPersons.objects.get(pk=person_id)
            old_name = f"{person.first_name} {person.surname}"
            old_email = person.email or ''
            
            person.first_name = self.request.POST.get('first_name', person.first_name)
            person.surname = self.request.POST.get('surname', person.surname)
            person.email = self.request.POST.get('email', person.email)
            
            new_name = f"{person.first_name} {person.surname}"
            new_email = person.email or ''
            
            changes = []
            if old_name != new_name:
                changes.append(f"Name changed from {old_name} to {new_name}")
            
            if old_email != new_email:
                changes.append(f"Email changed from {old_email} to {new_email}")
            
            if changes:
                change_note = f"{' and '.join(changes)} on {timezone.now().strftime('%Y-%m-%d')}"
                if person.comments:
                    person.comments += f"\n{change_note}"
                else:
                    person.comments = change_note
                        
            person.save()
            messages.success(self.request, "Successfully updated person information")
            
        except Exception as e:
            messages.error(self.request, f"Error updating person: {str(e)}")
        
        return self.get(self.request)
    

class TagRegisterView(LoginRequiredMixin, FormView):
    template_name = 'wamtram2/tag_register.html'
    form_class = TagRegisterForm
    success_url = reverse_lazy('wamtram2:tag_register')

    def dispatch(self, request, *args, **kwargs):
        if not (request.user.is_superuser or 
                request.user.groups.filter(name="WAMTRAM2_STAFF").exists()):
            raise PermissionDenied
        return super().dispatch(request, *args, **kwargs)

    def form_valid(self, form):
        try:
            tag_type = form.cleaned_data['tag_type']
            prefix = form.cleaned_data['tag_prefix']
            start = int(form.cleaned_data['start_number'])
            end = int(form.cleaned_data['end_number'])
            
            if end - start > 1000: 
                return JsonResponse({
                    'success': False, 
                    'error': 'Cannot create more than 1000 tags at once'
                })

            with transaction.atomic():  
                for num in range(start, end + 1):
                    
                    if tag_type == 'flipper':
                        tag_id = f"{prefix}{str(num).zfill(len(str(start)))}"
                    else:  # pit tags
                        tag_id = str(num)
                    
                    if tag_type == 'flipper':
                        
                        if TrtTags.objects.filter(tag_id=tag_id).exists():
                            return JsonResponse({
                                'success': False,
                                'error': f'Tag {tag_id} already exists'
                            })
                            
                        tag_status = TrtTagStatus.objects.get(tag_status='U')
                        
                        TrtTags.objects.create(
                            tag_id=tag_id,
                            tag_order_id=form.cleaned_data['tag_order_id'],
                            issue_location=form.cleaned_data['issue_location'],
                            custodian_person_id=form.cleaned_data['custodian_person_id'],
                            field_person_id=form.cleaned_data['field_person_id'],
                            comments=form.cleaned_data['comments'],
                            tag_status=tag_status
                        )
                    else:  # pit tags
                        
                        if TrtPitTags.objects.filter(pittag_id=tag_id).exists():
                            return JsonResponse({
                                'success': False,
                                'error': f'PIT tag {tag_id} already exists'
                            })
                            
                        pit_tag_status = TrtPitTagStatus.objects.get(pit_tag_status='U')
                            
                        TrtPitTags.objects.create(
                            pittag_id=tag_id,
                            tag_order_id=form.cleaned_data['tag_order_id'],
                            issue_location=form.cleaned_data['issue_location'],
                            custodian_person_id=form.cleaned_data['custodian_person_id'],
                            field_person_id=form.cleaned_data['field_person_id'],
                            comments=form.cleaned_data['comments'],
                            pit_tag_status=pit_tag_status
                        )

            return JsonResponse({
                'success': True,
                'message': f'Successfully registered {end - start + 1} tags'
            })

        except Exception as e:
            return JsonResponse({
                'success': False,
                'error': str(e)
            })
        
    def form_invalid(self, form):
        errors = []
        for field, error_list in form.errors.items():
            errors.append(f"{field}: {', '.join(error_list)}")
        
        return JsonResponse({
            'success': False,
            'error': 'Invalid form data: ' + '; '.join(errors)
        })


class AdminToolsView(LoginRequiredMixin, UserPassesTestMixin, TemplateView):
    template_name = 'wamtram2/admin_tools.html'
    
    def test_func(self):
        return (self.request.user.is_superuser)
    

class PitTagsListView(LoginRequiredMixin, UserPassesTestMixin, PaginateMixin, ListView):
    model = TrtPitTags
    template_name = 'wamtram2/pit_tags_list.html'
    context_object_name = 'pit_tags'
    paginate_by = 30
    
    def test_func(self):
        return self.request.user.is_superuser
    
    def get_queryset(self):
        queryset = super().get_queryset().select_related('turtle', 'custodian_person', 'pit_tag_status')
        
        search = self.request.GET.get('search')
        if search:
            queryset = queryset.filter(
                Q(pittag_id__icontains=search) |
                Q(turtle__turtle_id__icontains=search) |
                Q(custodian_person__first_name__icontains=search) |
                Q(custodian_person__surname__icontains=search)
            )
        
        status = self.request.GET.get('status')
        if status:
            queryset = queryset.filter(pit_tag_status=status)
            
        return queryset
        
    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        context.update({
            'admin_add_url': 'admin:wamtram2_trtpittags_add',
            'admin_change_url': 'admin:wamtram2_trtpittags_change',
            'search_term': self.request.GET.get('search', ''),
            'current_status': self.request.GET.get('status', ''),
            'status_choices': TrtPitTagStatus.objects.all(),
        })
        return context
    

class FlipperTagsListView(LoginRequiredMixin, UserPassesTestMixin, PaginateMixin, ListView):
    model = TrtTags
    template_name = 'wamtram2/flipper_tags_list.html'
    context_object_name = 'flipper_tags'
    paginate_by = 30
    
    def test_func(self):
        return self.request.user.is_superuser
    
    def get_queryset(self):
        queryset = super().get_queryset().select_related(
            'turtle', 
            'tag_status',
            'custodian_person'
        )
        
        search = self.request.GET.get('search')
        if search:
            queryset = queryset.filter(
                Q(tag_id__icontains=search) |
                Q(turtle__turtle_id__icontains=search) |
                Q(custodian_person__first_name__icontains=search) |
                Q(custodian_person__surname__icontains=search)
            )
        
        status = self.request.GET.get('status')
        if status:
            queryset = queryset.filter(tag_status_id=status)
            
        return queryset
        
    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        context.update({
            'admin_add_url': 'admin:wamtram2_trttags_add',
            'admin_change_url': 'admin:wamtram2_trttags_change',
            'search_term': self.request.GET.get('search', ''),
            'current_status': self.request.GET.get('status', ''),
            'status_choices': TrtTagStatus.objects.all(),
        })
        return context
            
class TransferObservationsByTagView(LoginRequiredMixin, View):
    template_name = 'wamtram2/transfer_observation.html'

    def dispatch(self, request, *args, **kwargs):
        if not (request.user.is_superuser or
                request.user.groups.filter(name="WAMTRAM2_STAFF").exists()):
            raise PermissionDenied
        return super().dispatch(request, *args, **kwargs)

    def get(self, request):
        return render(request, self.template_name)

    def get_turtle_info(self, turtle_id):
        """Get turtle information"""
        try:
            turtle = TrtTurtles.objects.get(turtle_id=turtle_id)
            return {
                'success': True,
                'data': {
                    'species': turtle.species_code.common_name,
                    'sex': turtle.sex,
                    'turtle_status': turtle.turtle_status.description,
                    'location_code': turtle.location_code.location_name,
                    'comments': turtle.comments
                }
            }
        except TrtTurtles.DoesNotExist:
            return {
                'success': False,
                'error': f'Turtle {turtle_id} not found'
            }

    def get_observations(self, tag_id):
        """Get observations data for a specific tag"""
        if not TrtTags.objects.filter(tag_id=tag_id).exists():
            return []
        
        observations = TrtObservations.objects.filter(
            trtrecordedtags__tag_id=tag_id
        ).select_related('turtle').values(
            'observation_id',
            'observation_date',
            'turtle_id',
            'place_code',
            'comments'
        ).order_by('-observation_date')
        return list(observations)

    def post(self, request):
        # Handle AJAX request for turtle info
        if request.headers.get('X-Requested-With') == 'FetchTurtleInfo':
            turtle_id = request.POST.get('turtle_id')
            return JsonResponse(self.get_turtle_info(turtle_id))

        # Handle AJAX request for observations
        if request.headers.get('X-Requested-With') == 'FetchObservations':
            tag_id = request.POST.get('tag_id')
            if not tag_id:
                return JsonResponse({
                    'success': False,
                    'error': 'Tag ID is required'
                })

            observations = self.get_observations(tag_id)
            return JsonResponse({
                'success': True,
                'observations': observations
            })

        # Handle transfer request  
        try:
            tag_id = request.POST.get('tag_id')
            turtle_id = request.POST.get('turtle_id')
            observation_ids = request.POST.getlist('observation_ids[]')
            
            # Basic validation
            if not all([tag_id, turtle_id]):
                return JsonResponse({
                    'success': False,
                    'error': 'Missing required parameters'
                }, status=400)
                
            # Convert observation_ids list to comma-separated string
            observation_ids_str = ','.join(observation_ids)

            # Execute stored procedure
            with connections['wamtram2'].cursor() as cursor:
                cursor.execute(
                    "EXEC dbo.TransferObservationsByFlipperTagWEB @TAG_ID = %s, @TURTLE_ID = %s, @OBSERVATION_IDS = %s;",
                    [tag_id, turtle_id, observation_ids_str]
                )
                
                # Get the results
                row = cursor.fetchone()
                return_value = row[0]
                error_message = row[1]

                if return_value == 0:
                    return JsonResponse({
                        'success': True,
                        'message': error_message
                    })
                else:
                    return JsonResponse({
                        'success': False,
                        'error': error_message
                    }, status=500)

        except Exception as e:
            return JsonResponse({
                'success': False,
                'error': str(e)
            }, status=500)
            
            
class NestingSeasonListView(LoginRequiredMixin, UserPassesTestMixin, PaginateMixin, ListView):
    model = TrtNestingSeason
    template_name = 'wamtram2/nesting_season_list.html'
    context_object_name = 'seasons'
    paginate_by = 30
    
    def test_func(self):
        return self.request.user.is_superuser
    
    def get_queryset(self):
        queryset = super().get_queryset()
        
        search = self.request.GET.get('search')
        if search:
            queryset = queryset.filter(
                Q(nesting_season__icontains=search)
            )
            
        return queryset
        
    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        context.update({
            'admin_add_url': reverse('admin:wamtram2_trtnestingseason_add'), 
            'admin_change_url': 'admin:wamtram2_trtnestingseason_change',
            'search_term': self.request.GET.get('search', ''),
        })
        return context
            
            
class BatchCurationView(LoginRequiredMixin, PaginateMixin,ListView):
    model = TrtEntryBatches
    template_name = 'wamtram2/batch_curation_list.html'
    context_object_name = 'batches'
    paginate_by = 30

    def dispatch(self, request, *args, **kwargs):
        if not (
            request.user.is_superuser
        ):
            return HttpResponseForbidden("You do not have permission to view this record")
        return super().dispatch(request, *args, **kwargs)

    def get_queryset(self):
        queryset = super().get_queryset().order_by('-entry_batch_id')
    
        queryset = queryset.select_related(
            'entered_person_id',
            'template'
        )

        queryset = queryset.annotate(
            entry_count=Count('trtdataentry'),
            flagged_count=Count('trtdataentry', filter=Q(trtdataentry__do_not_process=True))
        )
        
        search = self.request.GET.get('search')
        if search:
            queryset = queryset.filter(
                Q(batches_code__icontains=search) |
                Q(comments__icontains=search) |
                Q(entered_person_id__first_name__icontains=search) |
                Q(entered_person_id__surname__icontains=search)
            )
        return queryset
    
    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        
        # Get all possible columns
        all_columns = [
            {'field': 'entry_batch_id', 'title': 'Batch ID', 'visible': True},
            {'field': 'batches_code', 'title': 'Batch Code', 'visible': True},
            {'field': 'entry_date', 'title': 'Entry Date', 'visible': True},
            {'field': 'entered_person_id', 'title': 'Entered By', 'visible': True},
            {'field': 'entry_count', 'title': 'Entry Count', 'visible': True},
            {'field': 'flagged_count', 'title': 'Flagged Count', 'visible': True},
            {'field': 'template', 'title': 'Template', 'visible': False},
            {'field': 'comments', 'title': 'Comments', 'visible': False},
            {'field': 'pr_date_convention', 'title': 'PR Date Convention', 'visible': False},
        ]
        
        # Get column display settings from user settings or session
        user_columns = self.request.session.get('batch_grid_columns', [col['field'] for col in all_columns if col['visible']])
        
        context.update({
            'all_columns': all_columns,
            'visible_columns': user_columns,
            'search_term': self.request.GET.get('search', ''),
            'clear_url': self.request.path,
        })
        return context

    def post(self, request, *args, **kwargs):
        # Handle column display settings update
        if request.headers.get('X-Requested-With') == 'XMLHttpRequest':
            visible_columns = request.POST.getlist('columns[]')
            request.session['batch_grid_columns'] = visible_columns
            return JsonResponse({'status': 'success'})
        return HttpResponseBadRequest()


class EntryCurationView(LoginRequiredMixin, PaginateMixin, ListView):
    model = TrtDataEntry
    template_name = 'wamtram2/entry_curation_list.html'
    context_object_name = 'entries'
    paginate_by = 10

    def get_queryset(self):
        batch_id = self.kwargs.get('batch_id')
        queryset = super().get_queryset().filter(entry_batch_id=batch_id)
            
        queryset = queryset.select_related(
            'observation_id',
            'species_code',
            'place_code',
            'activity_code',
            'nesting',
            'interrupted',
            'alive',
            'measured_by_id',
            'recorded_by_id',
            'tagged_by_id',
            'entered_by_id',
            'measured_recorded_by_id',
            'egg_count_method',
            'clutch_completed',
            'flipper_tag_check',
            'pit_tag_check',
            'injury_check',
            'scar_check',
            'recapture_left_tag_id',
            'recapture_left_tag_id_2',
            'recapture_left_tag_id_3',
            'recapture_right_tag_id',
            'recapture_right_tag_id_2',
            'recapture_right_tag_id_3',
            'recapture_pittag_id',
            'recapture_pittag_id_2',
            'recapture_pittag_id_3',
            'recapture_pittag_id_4',
            'new_left_tag_id',
            'new_left_tag_id_2',
            'new_right_tag_id',
            'new_right_tag_id_2',
            'new_pittag_id',
            'new_pittag_id_2',
            'new_pittag_id_3',
            'new_pittag_id_4',
            'body_part_1',
            'body_part_2',
            'body_part_3',
            'damage_code_1',
            'damage_code_2',
            'damage_code_3',
            'tissue_type_1',
            'tissue_type_2',
            'measurement_type_1',
            'measurement_type_2',
            'measurement_type_3',
            'measurement_type_4',
            'measurement_type_5',
            'measurement_type_6',
            'recapture_left_tag_state',
            'recapture_left_tag_state_2',
            'recapture_right_tag_state',
            'recapture_right_tag_state_2',
            'new_left_tag_state',
            'new_left_tag_state_2',
            'new_right_tag_state',
            'new_right_tag_state_2'
        )
        
        filter_value = self.request.GET.get("filter")
        if filter_value == "needs_review":
            queryset = queryset.filter(do_not_process=True)
        elif filter_value == "not_saved":
            queryset = queryset.filter(observation_id__isnull=True)
        elif filter_value == "needs_review_no_message":
            queryset = queryset.filter(do_not_process=True, error_message__isnull=True)

        search = self.request.GET.get('search')
        if search:
            queryset = queryset.filter(
                Q(comments__icontains=search) |
                Q(turtle_comments__icontains=search) |
                Q(species_code__code__icontains=search)
            )
        if not queryset.exists():
            return TrtDataEntry.objects.none()
        
        return queryset.order_by("-data_entry_id")
        
    def get_context_data(self, **kwargs):
            context = super().get_context_data(**kwargs)
            context['sex_choices'] = TrtDataEntry.SEX_CHOICES
            if not context.get('object_list'):
                context['object_list'] = []
            
            context.update({
                'species_choices': TrtSpecies.objects.all(),
                'places_choices': TrtPlaces.objects.select_related('location_code').all(),
                'activities_choices': TrtActivities.objects.all(),
                'yesno_choices': TrtYesNo.objects.all(),
                'persons_choices': TrtPersons.objects.all(),
                'damage_codes_choices': TrtDamageCodes.objects.all(),
                'body_parts_choices': TrtBodyParts.objects.all(),
                'measurement_types_choices': TrtMeasurementTypes.objects.all(),
                'tag_states_choices': TrtTagStates.objects.all(),
                'tissue_types_choices': TrtTissueTypes.objects.all(),
                'egg_count_methods_choices': TrtEggCountMethods.objects.all(),
                'identification_types_choices': TrtIdentificationTypes.objects.all(),
                'tags_choices': TrtTags.objects.all(),
                'pit_tags_choices': TrtPitTags.objects.all(),
            })
            
            model_fields = TrtDataEntry._meta.get_fields()
    
            default_visible_fields = {
                'data_entry_id','observation_id', 'turtle_id', 'entered_by_id', 'species_code', 'place_code',
                'observation_date', 'do_not_process','error_message', 'comments'
            }
            
            field_groups = {
                'Basic Information': [
                    'data_entry_id', 'observation_id', 'turtle_id', 'comments','error_message','species_code', 'sex', 'place_code', 
                    'observation_date', 'observation_time', 'do_not_process', 'latitude', 'longitude',
                ],
                'Measurements': [
                    'curved_carapace_length', 'curved_carapace_width',
                    'curved_carapace_length_notch', 'cc_length_not_measured',
                    'cc_width_not_measured', 'cc_notch_length_not_measured',
                    'measurement_type_1', 'measurement_value_1',
                    'measurement_type_2', 'measurement_value_2',
                    'measurement_type_3', 'measurement_value_3',
                    'measurement_type_4', 'measurement_value_4',
                    'measurement_type_5', 'measurement_value_5',
                    'measurement_type_6', 'measurement_value_6'
                ],
                'Status and Activities': [
                    'alive', 'activity_code', 'nesting', 'interrupted',
                    'identification_confidence'
                ],
                'Flipper Tags': [
                    'flipper_tag_check',
                    'recapture_left_tag_id', 'recapture_left_tag_state', 'recapture_left_tag_position', 'recapture_left_tag_state', 'recapture_left_tag_barnacles'
                    'recapture_left_tag_id_2', 'recapture_left_tag_state_2', 'recapture_left_tag_position_2', 'recapture_left_tag_state_2', 'recapture_left_tag_barnacles_2',
                    'recapture_right_tag_id', 'recapture_right_tag_state', 'recapture_right_tag_position', 'recapture_right_tag_state', 'recapture_right_tag_barnacles',
                    'recapture_right_tag_id_2', 'recapture_right_tag_state_2', 'recapture_right_tag_position_2', 'recapture_right_tag_state_2', 'recapture_right_tag_barnacles_2',

                    'new_left_tag_id', 'new_left_tag_state', 'new_left_tag_position',
                    'new_left_tag_id_2', 'new_left_tag_state_2', 'new_left_tag_position_2',
                    'new_right_tag_id', 'new_right_tag_state', 'new_right_tag_position',
                    'new_right_tag_id_2', 'new_right_tag_state_2', 'new_right_tag_position_2',

                    'other_left_tag', 
                    'other_right_tag', 
                ],
                'PIT Tags': [
                    'pit_tag_check',
                    'recapture_pittag_id', 
                    'recapture_pittag_id_2', 
                    'recapture_pittag_id_3', 
                    'recapture_pittag_id_4',
                    'new_pittag_id', 'new_pit_tag_sticker_present',
                    'new_pittag_id_2', 'new_pit_tag_2_sticker_present',
                    'new_pittag_id_3', 'new_pit_tag_3_sticker_present', 
                    'new_pittag_id_4','new_pit_tag_4_sticker_present',
                ],
                'Dud Tags': [
                    'dud_flipper_tag', 'dud_flipper_tag_2', 'dud_pit_tag', 'dud_pit_tag_2'
                ],
                'Other Tags': [
                    'other_tags', 'other_tags_identification_type'
                    'identifer', 'identification_type'
                ],
                'Tag Scars': [
                    'scar_check',
                    'scars_left', 'scars_right',
                    'scars_left_scale_1', 'scars_left_scale_2', 'scars_left_scale_3',
                    'scars_right_scale_1', 'scars_right_scale_2', 'scars_right_scale_3'
                ],
                'Damage and Scars': [
                    'injury_check',
                    'body_part_1', 'damage_code_1',
                    'body_part_2', 'damage_code_2',
                    'body_part_3', 'damage_code_3',
                    'body_part_4', 'damage_code_4',
                    'body_part_5', 'damage_code_5',
                    'body_part_6', 'damage_code_6',
                ],
                'Tissue Samples': [
                    'tissue_type_1', 'sample_label_1',
                    'tissue_type_2', 'sample_label_2'
                ],
                'Nesting Data': [
                    'egg_count', 'egg_count_method', 'clutch_completed'
                ],
                'Personnel': [
                    'entered_by_id',
                    'measured_by_id',
                    'recorded_by_id',
                    'tagged_by_id',
                ],
                'Comments': [
                    'turtle_comments', 'comment_fromrecordedtagstable',
                    'error_number'
                ]
            }
    
    
            field_order = {}
            order_index = 0
            for group_fields in field_groups.values():
                for field in group_fields:
                    field_order[field] = order_index
                    order_index += 1
    
            all_columns = []
            processed_fields = set()
    
    
            for group_name, field_patterns in field_groups.items():
                group_fields = []
                for field in model_fields:
                    if hasattr(field, 'name'):
                        field_name = field.name
                        if field_name in field_patterns:
                            group_fields.append({
                                'field': field_name,
                                'title': field_name.replace('_', ' ').title(),
                                'visible': field_name in default_visible_fields,
                                'group': group_name
                            })
                            processed_fields.add(field_name)
                
                sorted_fields = sorted(group_fields, key=lambda x: field_order.get(x['field'], float('inf')))
                all_columns.extend(sorted_fields)
    
    
            other_fields = []
            for field in model_fields:
                if hasattr(field, 'name') and field.name not in processed_fields:
                    other_fields.append({
                        'field': field.name,
                        'title': field.name.replace('_', ' ').title(),
                        'visible': field.name in default_visible_fields,
                        'group': 'Other'
                    })
            
    
            all_columns.extend(sorted(other_fields, key=lambda x: x['field']))
    
    
            user_columns = self.request.session.get('entry_grid_columns',
                                                [col['field'] for col in all_columns if col['visible']])
            
            batch_id = self.kwargs.get('batch_id')
            
    
            context.update({
                'all_columns': all_columns,
                'visible_columns': user_columns,
                'search_term': self.request.GET.get('search', ''),
                'clear_url': reverse('wamtram2:entries_curation', kwargs={'batch_id': batch_id}),
                'batch_id': batch_id,
            })
            
            return context
    
    def post(self, request, *args, **kwargs):
        if request.headers.get('X-Requested-With') == 'XMLHttpRequest':
            visible_columns = request.POST.getlist('columns[]')
            request.session['entry_grid_columns'] = visible_columns
            return JsonResponse({'status': 'success'})
        return HttpResponseBadRequest()


class SaveEntryChangesView(LoginRequiredMixin, View):
    READONLY_FIELDS = {'data_entry_id', 'observation_id'}
    
    def validate_field(self, field_name, value, entry):
        if field_name in self.READONLY_FIELDS:
            raise ValueError(f"Field {field_name} is readonly")
            
        field = TrtDataEntry._meta.get_field(field_name)
        
        if field.is_relation:
            if value == '':
                return None
            try:
                related_model = field.related_model
                pk_name = related_model._meta.pk.name
                lookup = {pk_name: value}
                instance = related_model.objects.get(**lookup)

                if field_name.endswith('_by_id'):
                    base_field = field_name[:-3]  
                    if hasattr(entry, base_field):
                        setattr(entry, base_field, f"{instance.first_name} {instance.surname}")
                return instance
            except related_model.DoesNotExist:
                raise ValueError(f"Invalid value for {field_name}: {value}")
        
 
        if field.get_internal_type() in ['IntegerField', 'FloatField']:
            try:
                value = float(value)
                if value < 0:
                    raise ValueError(f"{field_name} cannot be negative")
                return value
            except ValueError:
                raise ValueError(f"Invalid number for {field_name}: {value}")
        
        if field.get_internal_type() == 'CharField':
            if not isinstance(value, str):
                raise ValueError(f"Invalid string for {field_name}: {value}")
            return value
        
        return value

    def post(self, request):
        try:
            changes = json.loads(request.POST.get('changes', '{}'))
            
            with transaction.atomic():
                for entry_id, fields in changes.items():
                    entry = TrtDataEntry.objects.get(data_entry_id=entry_id)
                    
                    for field, value in fields.items():
                        validated_value = self.validate_field(field, value, entry)
                        setattr(entry, field, validated_value)
                    
                    entry.save()
                    
            return JsonResponse({'success': True})
            
        except Exception as e:
            return JsonResponse({
                'success': False,
                'error': str(e)
            })


class ObservationManagementView(LoginRequiredMixin, TemplateView):
    template_name = 'wamtram2/observation_management.html'
    
    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        observation_id = self.kwargs.get('observation_id')
        
        context['initial_data'] = 'null'
        if observation_id:
            try:
                print(f"尝试获取观察记录数据，ID: {observation_id}")  # 添加日志
                observation_data_view = ObservationDataView()
                response = observation_data_view.get(self.request, observation_id)
                print(f"响应状态码: {response.status_code}")  # 添加日志
                print(f"响应内容: {response.content}")  # 添加日志
                if response.status_code == 200:
                    data = json.loads(response.content)
                    if data['status'] == 'success':
                        context['initial_data'] = json.dumps(data['data'])
                        print(f"成功设置 initial_data: {context['initial_data']}")  # 添加日志
                    else:
                        print(f"响应状态不是success: {data['status']}")  # 添加日志
                else:
                    print(f"响应状态码不是200: {response.status_code}")  # 添加日志
            except Exception as e:
                print(f"获取观察记录数据时出错: {str(e)}")  # 添加日志
                import traceback
                print(traceback.format_exc())  # 打印完整的错误堆栈
        # 序列化选项数据
        try:
            context.update({
                'tag_states_choices': [
                    {
                        'tag_state': state.tag_state,
                        'description': state.description
                    } for state in TrtTagStates.objects.all()
                ],
                'pit_tag_state_choices': [
                    {
                        'pit_tag_state': state.pit_tag_state,
                        'description': state.description
                    } for state in TrtPitTagStates.objects.all()
                ],
                'measurement_type_choices': [
                    {
                        'measurement_type': type.measurement_type,
                        'description': type.description
                    } for type in TrtMeasurementTypes.objects.all()
                ],
                'body_parts_choices': [
                    {
                        'body_part': part.body_part,
                        'description': part.description
                    } for part in TrtBodyParts.objects.all()
                ],
                'damage_codes_choices': [
                    {
                        'damage_code': code.damage_code,
                        'description': code.description
                    } for code in TrtDamageCodes.objects.all()
                ],
                'damage_cause_choices': [
                    {
                        'damage_cause_code': cause.damage_cause_code,
                        'description': cause.description
                    } for cause in TrtDamageCauseCodes.objects.all()
                ],
                'identification_type_choices': [
                    {
                        'identification_type': type.identification_type,
                        'description': type.description
                    } for type in TrtIdentificationTypes.objects.all()
                ],
                'date_convention_choices': [
                    {'code': 'C', 'description': 'Calendar'},
                    {'code': 'E', 'description': 'Evening'},
                    {'code': 'U', 'description': 'Unknown'}
                ],
                'places': TrtPlaces.objects.all(),
                'activity_code_choices': TrtActivities.objects.all(),
                'beach_position_code_choices': TrtBeachPositions.objects.all(),
                'yes_no_choices': TrtYesNo.objects.all(),
                'datum_code_choices': TrtDatumCodes.objects.all(),
                'condition_code_choices': TrtConditionCodes.objects.all(),
                'egg_count_method_choices': TrtEggCountMethods.objects.all(),
                'search_persons_url': reverse('wamtram2:search-persons'),
                'search_places_url': reverse('wamtram2:search-places'),
                'submit_url': reverse('wamtram2:observation_detail', kwargs={'observation_id': observation_id}) if observation_id else reverse('wamtram2:observation_detail'),
            })
        except Exception as e:
            print(f"更新context时出错: {str(e)}") 
            import traceback
            print(traceback.format_exc()) 
            
            for key in ['tag_states_choices', 'pit_tag_state_choices', 'measurement_type_choices',
                'body_parts_choices', 'damage_codes_choices', 'damage_cause_choices']:
                context[key] = json.dumps(context[key])
        return context


class ObservationDataView(LoginRequiredMixin, View):
    @transaction.atomic
    def post(self, request, observation_id=None): 
        try:
            data = json.loads(request.body)
            
            observation_id = observation_id or data.get('observation_id')
            if observation_id:
                observation = TrtObservations.objects.get(pk=observation_id)
            else:
                observation = TrtObservations()
            
            # Update basic info
            basic_info = data.get('basic_info', {})
            for field, value in basic_info.items():
                if hasattr(observation, field):
                    setattr(observation, field, value)
            
            if 'observation_date' in basic_info:
                try:
                    datetime_obj = datetime.strptime(
                        basic_info['observation_date'], 
                        '%Y-%m-%dT%H:%M'
                    )
                    observation.observation_date = datetime_obj
                    observation.observation_time = datetime_obj
                except ValueError as e:
                    raise ValidationError(f"Invalid date format: {str(e)}")
            
            observation.save()
            
            # Update related records
            self._update_tags(observation, data.get('tag_info', {}))
            self._update_measurements(observation, data.get('measurements', []))
            self._update_damage_records(observation, data.get('damage_records', []))
            self._update_identifications(observation, data.get('recorded_identifications', [])) 
            self._update_location(observation, data.get('location', {}))
            
            return JsonResponse({'status': 'success', 'observation_id': observation.observation_id})
            
        except Exception as e:
            return JsonResponse({'status': 'error', 'message': str(e)}, 
                            status=400 if isinstance(e, ValidationError) else 500)

    def get(self, request, observation_id=None):
        try:
            if observation_id:
                print(f"尝试获取观察记录 ID: {observation_id}")  # 添加日志
                observation = TrtObservations.objects.get(pk=observation_id)
                try:
                    data = self._get_observation_data(observation)
                    print(f"成功获取观察数据: {data}")  # 添加日志
                    return JsonResponse({
                        'status': 'success', 
                        'data': data
                    })
                except Exception as e:
                    print(f"获取观察数据时出错: {str(e)}")  # 添加日志
                    import traceback
                    print(traceback.format_exc())  # 打印完整的错误堆栈
                    return JsonResponse({
                        'status': 'error',
                        'message': f'Error getting observation data: {str(e)}'
                    }, status=500)
            else:
                observations = self._filter_observations(request)
                data = [self._get_observation_summary(obs) for obs in observations]
                return JsonResponse({'status': 'success', 'data': data})
        except TrtObservations.DoesNotExist:
            print(f"未找到观察记录 ID: {observation_id}")  # 添加日志
            return JsonResponse({
                'status': 'error', 
                'message': f'Observation {observation_id} not found'
            }, status=404)
        except Exception as e:
            print(f"处理请求时出错: {str(e)}")  # 添加日志
            import traceback
            print(traceback.format_exc())  # 打印完整的错误堆栈
            return JsonResponse({
                'status': 'error',
                'message': f'Error processing request: {str(e)}'
            }, status=500)
    
    def _get_observation_data(self, observation):
        """Get full observation data"""
        scars_data = {
            'scars_left': observation.scars_left,
            'scars_right': observation.scars_right,
            'scars_left_scale_1': observation.scars_left_scale_1,
            'scars_left_scale_2': observation.scars_left_scale_2,
            'scars_left_scale_3': observation.scars_left_scale_3,
            'scars_right_scale_1': observation.scars_right_scale_1,
            'scars_right_scale_2': observation.scars_right_scale_2,
            'scars_right_scale_3': observation.scars_right_scale_3,
            'tag_scar_not_checked': observation.tagscarnotchecked
        }
        damage_data = [{
            'body_part': damage.body_part.body_part if damage.body_part else None,
            'damage_code': damage.damage_code.damage_code if damage.damage_code else None,
            'damage_cause_code': damage.damage_cause_code.damage_cause_code if damage.damage_cause_code else None,
            'comments': damage.comments
        } for damage in TrtDamage.objects.filter(observation=observation)]
        
        other_tags_data = {
            'other_tags': observation.other_tags,
            'identification_type': observation.other_tags_identification_type.identification_type if observation.other_tags_identification_type else None,
            'identification_type_description': observation.other_tags_identification_type.description if observation.other_tags_identification_type else None
        }
            
        recorded_identifications = []
        for record in TrtRecordedIdentification.objects.filter(observation_id=observation.observation_id):
            try:
                # 获取所有关联的记录
                turtles = TrtIdentification.objects.filter(turtle2=record)
                for turtle in turtles:
                    recorded_identifications.append({
                        'turtle_id': turtle.turtle_id,
                        'identification_type': record.identification_type.identification_type if record.identification_type else None,
                        'identifier': record.identifier.identifier if record.identifier else None,
                        'comments': record.comments
                    })
                    
            except Exception as e:
                print(f"Error processing recorded identification: {str(e)}")
                continue
                
        persons_data = {
            'measurer_person': {
                'id': observation.measurer_person.person_id if observation.measurer_person else None,
                'text': str(observation.measurer_person) if observation.measurer_person else None
            },
            'measurer_reporter_person': {
                'id': observation.measurer_reporter_person.person_id if observation.measurer_reporter_person else None,
                'text': str(observation.measurer_reporter_person) if observation.measurer_reporter_person else None
            },
            'tagger_person': {
                'id': observation.tagger_person.person_id if observation.tagger_person else None,
                'text': str(observation.tagger_person) if observation.tagger_person else None
            },
            'reporter_person': {
                'id': observation.reporter_person.person_id if observation.reporter_person else None,
                'text': str(observation.reporter_person) if observation.reporter_person else None
            }
        }
        place_data = None
        if observation.place_code:
            place_data = {
                'id': observation.place_code.place_code if observation.place_code else None,
                'text': observation.place_code.get_full_name() if observation.place_code else None
            }
    
        tag_info = {
            'recorded_tags': [{
                'tag_id': str(tag.tag_id),
                'tag_side': tag.side,
                'tag_position': tag.tag_position,
                'tag_state': tag.tag_state.tag_state if tag.tag_state else None,
                'comments': tag.comments
            } for tag in observation.trtrecordedtags_set.all()],
            'recorded_pit_tags': [{
                'tag_id': str(tag.pittag_id),
                'tag_position': tag.pit_tag_position,
                'tag_state': tag.pit_tag_state.pit_tag_state if tag.pit_tag_state else None,
                'comments': tag.comments
            } for tag in observation.trtrecordedpittags_set.all()]
        }
    
        measurements = [{
            'measurement_type': str(measurement.measurement_type.measurement_type) if measurement.measurement_type else None,
            'measurement_value': str(measurement.measurement_value),
            'comments': measurement.comments
        } for measurement in observation.trtmeasurements_set.all()]
    
        return {
            'basic_info': {
                'observation_id': observation.observation_id,
                'turtle_id': observation.turtle_id,
                'observation_date': observation.observation_date.strftime('%Y-%m-%dT%H:%M') if observation.observation_date else '', 
                'alive': str(observation.alive.code) if observation.alive else '',
                'nesting': str(observation.nesting.code) if observation.nesting else '',
                'activity_code': str(observation.activity_code.activity_code) if observation.activity_code else '',
                'beach_position_code': str(observation.beach_position_code.beach_position_code) if observation.beach_position_code else '',
                'condition_code': str(observation.condition_code.condition_code) if observation.condition_code else '',
                'number_of_eggs': str(observation.number_of_eggs) if observation.number_of_eggs else '',
                'egg_count_method': str(observation.egg_count_method.egg_count_method) if observation.egg_count_method else '',
                'observation_status': str(observation.observation_status),
                'measurer_person': persons_data['measurer_person'],
                'measurer_reporter_person': persons_data['measurer_reporter_person'],
                'tagger_person': persons_data['tagger_person'],
                'reporter_person': persons_data['reporter_person'],
                'place_code': place_data,
                'datum_code': str(observation.datum_code) if observation.datum_code else '',
                'latitude': str(observation.latitude) if observation.latitude else '',
                'longitude': str(observation.longitude) if observation.longitude else '',
                'clutch_completed': str(observation.clutch_completed) if observation.clutch_completed else '',
                'date_convention': str(observation.date_convention) if observation.date_convention else '',
            },
            'tag_info': tag_info,
            'measurements': measurements,
            'damage_records': damage_data,
            'recorded_identifications': recorded_identifications,
            'other_tags_data': other_tags_data,
            'scars': scars_data
        }
        
    def _filter_observations(self, request):
        """Filter observations based on request parameters"""
        observations = TrtObservations.objects.all()
        search_term = request.GET.get('search')
        if search_term:
            tag_parts = search_term.split()
            q_objects = Q()
            for part in tag_parts:
                q_objects |= (
                    Q(trtrecordedtags__tag_id__tag_id__exact=part) | 
                    Q(trtrecordedpittags__pittag_id__pittag_id__exact=part) 
                )
            observations = observations.filter(q_objects)
        
        place_code = request.GET.get('place')
        if place_code:
            observations = observations.filter(place_code=place_code)
        
        date = request.GET.get('date')
        if date:
            try:
                date_obj = datetime.strptime(date, '%Y-%m-%d')
                observations = observations.filter(observation_date__date=date_obj)
            except ValueError:
                pass
        
        status = request.GET.get('status')
        if status:
            if status == 'Initial Nesting':
                observations = observations.filter(nesting='Y', status='I')
            elif status == 'Subsequent Nesting':
                observations = observations.filter(nesting='Y', status='S')
            elif status == 'Non-nesting':
                observations = observations.filter(nesting='N')
                
        observations = observations.select_related('place_code').prefetch_related(
            'trtrecordedtags_set', 
            'trtrecordedpittags_set'
            ).distinct().order_by('-observation_date', '-observation_time')
        
        return observations.order_by('-observation_date')
    
    def _get_observation_summary(self, observation):
        """Get summary data for observation list"""
        tags = list(observation.trtrecordedtags_set.values_list('tag_id', flat=True)) 
        pit_tags = list(observation.trtrecordedpittags_set.values_list('pittag_id', flat=True))
        place = observation.place_code
        place_description = place.get_full_name() if place else ''
        return {
            'observation_id': observation.observation_id,
            'turtle_id': observation.turtle_id,
            'observation_date': observation.observation_date.strftime('%Y-%m-%d'),
            'observation_time': observation.observation_time.strftime('%H:%M') if observation.observation_time else '',
            'place_code': place.place_code if place else '',
            'place_description': place_description,
            'status': observation.observation_status,
            'comments': observation.comments,
            'nesting': str(observation.nesting),
            'tags': [str(tag) for tag in tags],
            'pit_tags': [str(tag) for tag in pit_tags],
            'total_tags': len(tags),
            'total_pit_tags': len(pit_tags)
        }


class SaveObservationView(LoginRequiredMixin, View):
    @transaction.atomic
    def post(self, request, observation_id=None):
        try:
            data = json.loads(request.body)
            
            observation_id = observation_id or data.get('observation_id')
            if observation_id:
                observation = TrtObservations.objects.get(pk=observation_id)
            else:
                observation = TrtObservations()
            

            self._update_basic_info(observation, data.get('basic_info', {}))
            observation.save()
            self._update_tags(observation, data.get('tag_info', {}))
            self._update_measurements(observation, data.get('measurements', []))
            self._update_damage_records(observation, data.get('damage_records', []))
            self._update_identifications(observation, data.get('recorded_identifications', []))
            self._update_location(observation, data.get('location', {}))
            self._update_scars(observation, data.get('scars', {}))
            self._update_other_tags(observation, data.get('other_tags_data', {}))
            self._update_status(observation)  
            
            return JsonResponse({
                'status': 'success', 
                'observation_id': observation.observation_id
            })
            
        except ValidationError as e:
            return JsonResponse({
                'status': 'error',
                'message': str(e)
            }, status=400)
        except Exception as e:
            print(f"Error during save record: {str(e)}")
            import traceback
            print(traceback.format_exc())
            return JsonResponse({
                'status': 'error',
                'message': f'Error during save record: {str(e)}'
            }, status=500)

    FOREIGN_KEY_FIELDS = {
        'other_tags_identification_type': TrtIdentificationTypes,
        'measurer_person': TrtPersons,
        'measurer_reporter_person': TrtPersons,
        'tagger_person': TrtPersons,
        'reporter_person': TrtPersons,
        'place_code': TrtPlaces,
        'alive': TrtYesNo,
        'nesting': TrtYesNo,
        'clutch_completed': TrtYesNo,
        'activity_code': TrtActivities,
        'beach_position_code': TrtBeachPositions,
        'condition_code': TrtConditionCodes,
        'egg_count_method': TrtEggCountMethods,
        'datum_code': TrtDatumCodes
    }


    def _update_basic_info(self, observation, basic_info):
        """Update basic information"""
        try:
            if 'observation_date' in basic_info:
                try:
                    datetime_obj = datetime.strptime(
                        basic_info['observation_date'], 
                        '%Y-%m-%dT%H:%M'
                    )
                    if timezone.is_naive(datetime_obj):
                        datetime_obj = timezone.make_aware(
                            datetime_obj, 
                            timezone.get_current_timezone()
                        )
                    observation.observation_date = datetime_obj
                    observation.observation_time = datetime_obj
                except ValueError as e:
                    raise ValidationError(f"Invalid date format: {str(e)}")

            # Update other basic fields
            for field, value in basic_info.items():
                if hasattr(observation, field):
                    # Check if it's a foreign key field
                    if field in self.FOREIGN_KEY_FIELDS:
                        if value:  # If there's a value
                            if isinstance(value, dict) and 'id' in value:  # Handle Select2 format
                                value = value['id']
                            try:
                                # Get the foreign key object
                                model_class = self.FOREIGN_KEY_FIELDS[field]
                                related_obj = model_class.objects.get(pk=value)
                                setattr(observation, field, related_obj)
                            except model_class.DoesNotExist:
                                print(f"Can't find the record for {field}: {value}")
                                setattr(observation, field, None)
                        else:  # If the value is empty, set to None
                            setattr(observation, field, None)
                    else:  # Non-foreign key fields are assigned directly
                        if field == 'number_of_eggs':
                            if value == '' or value is None:
                                value = None
                            else:
                                try:
                                    value = int(value)
                                except ValueError:
                                    value = None
                        setattr(observation, field, value)

        except Exception as e:
            print(f"Error updating basic information: {str(e)}")
            import traceback
            print(traceback.format_exc())
            raise ValidationError(f"Error updating basic information: {str(e)}")
        
    def _update_tags(self, observation, tag_data):
        """Update tag records tag records tag records"""
        try:    
            # Get existing recordsisting recordsisting records
            existing_tags = {
                tag.tag_id_id: tag 
                for tag in TrtRecordedTags.objects.filter(observation_id=observation)
            }
            existing_pit_tags = {
                pit.pittag_id_id: pit 
                for pit in TrtRecordedPitTags.objects.filter(observation_id=observation)
            }
            
            # Process normal tags
            processed_tag_ids = set()
            for tag in tag_data.get('recorded_tags', []):
                if tag.get('tag_id'):
                    try:
                        tag_instance = TrtTags.objects.get(tag_id=tag['tag_id'])
                        processed_tag_ids.add(tag_instance.tag_id)
                        
                        if tag_instance.tag_id in existing_tags:
                            # Update existing record
                            existing_tag = existing_tags[tag_instance.tag_id]
                            existing_tag.side = tag.get('tag_side')
                            existing_tag.tag_position = tag.get('tag_position')
                            existing_tag.tag_state_id = tag.get('tag_state')
                            existing_tag.barnacles = False
                            existing_tag.turtle_id = observation.turtle_id
                            existing_tag.save()
                        else:
                            TrtRecordedTags.objects.create(
                                observation_id=observation,
                                tag_id=tag_instance,
                                side=tag.get('tag_side'),
                                tag_position=tag.get('tag_position'),
                                tag_state_id=tag.get('tag_state'),
                                barnacles=False,
                                turtle_id=observation.turtle_id
                            )
                    except TrtTags.DoesNotExist:
                        print(f"找不到标签记录: {tag['tag_id']}")
                        continue
            
            # Process PIT tags
            processed_pit_tag_ids = set()
            for pit_tag in tag_data.get('recorded_pit_tags', []):
                if pit_tag.get('tag_id'):
                    try:
                        pit_tag_instance = TrtPitTags.objects.get(pittag_id=pit_tag['tag_id'])
                        processed_pit_tag_ids.add(pit_tag_instance.pittag_id)
                        
                        if pit_tag_instance.pittag_id in existing_pit_tags:
                            # Update existing record
                            existing_pit = existing_pit_tags[pit_tag_instance.pittag_id]
                            existing_pit.pit_tag_position = pit_tag.get('tag_position')
                            existing_pit.pit_tag_state_id = pit_tag.get('tag_state')
                            existing_pit.save()
                        else:
                            # Create new record
                            TrtRecordedPitTags.objects.create(
                                observation_id=observation,
                                pittag_id=pit_tag_instance,
                                pit_tag_position=pit_tag.get('tag_position'),
                                pit_tag_state_id=pit_tag.get('tag_state'),
                                turtle_id=observation.turtle,
                                checked=True
                            )
                    except TrtPitTags.DoesNotExist:
                        print(f"Can't find the record for PIT tag: {pit_tag['tag_id']}")
                        continue

        except Exception as e:
            print(f"Error updating tag records: {str(e)}")
            print(traceback.format_exc())
            raise ValidationError(f"Error updating tag records: {str(e)}")
        
    def _update_location(self, observation, location_data):
        """Update location information"""
        if location_data:
            for field in ['place_code', 'datum_code', 'latitude', 'longitude']:
                if field in location_data:
                    if field in self.FOREIGN_KEY_FIELDS and location_data[field]:
                        try:
                            related_obj = self.FOREIGN_KEY_FIELDS[field].objects.get(
                                pk=location_data[field]
                            )
                            setattr(observation, field, related_obj)
                        except self.FOREIGN_KEY_FIELDS[field].DoesNotExist:
                            setattr(observation, field, None)
                    else:
                        # Handle empty values for numeric fields
                        value = location_data[field]
                        if field in ['latitude', 'longitude']:
                            # If it's an empty string or None, set to None
                            if value == '' or value is None:
                                value = None
                            # If it's a valid numeric string, convert to float
                            elif isinstance(value, str):
                                try:
                                    value = float(value)
                                except ValueError:
                                    value = None
                        setattr(observation, field, value)
            observation.save()

    def _update_identifications(self, observation, identification_data):
        """Update identification records"""
        try:
            # Get existing records
            existing_identifications = {
                str(record.turtle_id): record 
                for record in TrtRecordedIdentification.objects.filter(
                    observation_id=observation.observation_id
                )
            }
            
            # Process new records
            for record in identification_data:
                if record.get('turtle_id'):
                    try:
                        # Process turtle_id
                        turtle_id = record['turtle_id']
                        if isinstance(turtle_id, str) and turtle_id.startswith('T'):
                            # If it starts with 'T', remove 'T'
                            numeric_turtle_id = turtle_id[1:]
                        else:
                            numeric_turtle_id = str(turtle_id)
                        
                        # Check if the record exists
                        if numeric_turtle_id in existing_identifications:
                            # Update existing record
                            existing_record = existing_identifications[numeric_turtle_id]
                            existing_record.identification_type_id = record.get('identification_type')
                            existing_record.identifier = record.get('identifier')
                            existing_record.comments = record.get('comments')
                            existing_record.save()
                        else:
                            # Create new record
                            TrtRecordedIdentification.objects.create(
                                observation_id=observation.observation_id,
                                turtle_id=numeric_turtle_id,
                                identification_type_id=record.get('identification_type'),
                                identifier=record.get('identifier'),
                                comments=record.get('comments')
                            )
                            
                    except (ValueError, TypeError) as e:
                        print(f"Error processing turtle_id: {str(e)}, turtle_id: {record.get('turtle_id')}")
                        continue
                    except Exception as e:
                        print(f"Error processing identification records: {str(e)}")
                        print(traceback.format_exc())
                        continue
                        
        except Exception as e:
            print(f"Error updating identification records: {str(e)}")
            print(traceback.format_exc())
            raise ValidationError(f"Error updating identification records: {str(e)}")

    def _update_measurements(self, observation, measurements):
        """Update measurement records"""
        try:
            # Get existing records
            existing_measurements = {
                measure.measurement_type_id: measure 
                for measure in observation.trtmeasurements_set.all()
            }
            
            #   
            for measurement in measurements:
                if measurement.get('measurement_value') is not None:  # Ensure there's a measurement value
                    measurement_type = measurement.get('measurement_type')
                    
                    if measurement_type in existing_measurements:
                        # Update existing record
                        existing_measure = existing_measurements[measurement_type]
                        existing_measure.measurement_value = measurement['measurement_value']
                        existing_measure.comments = measurement.get('comments')
                        existing_measure.save()
                    else:
                        # Create new record
                        TrtMeasurements.objects.create(
                            observation_id=observation,
                            measurement_type_id=measurement_type,
                            measurement_value=measurement['measurement_value'],
                            comments=measurement.get('comments')
                        )

        except Exception as e:
            print(f"Error updating measurement records: {str(e)}")
            print(traceback.format_exc())
            raise ValidationError(f"Error updating measurement records: {str(e)}")

    def _update_damage_records(self, observation, damage_records):
        """Update damage records"""
        try:
            # Get existing records
            existing_damages = {
                damage.body_part_id: damage 
                for damage in observation.damages.all()
            }
            
            # Process each damage record
            for damage in damage_records:
                if damage.get('body_part') and damage.get('damage_code'):
                    body_part_id = damage['body_part']
                    
                    try:
                        if body_part_id in existing_damages:
                            # Update existing record
                            existing_damage = existing_damages[body_part_id]
                            existing_damage.damage_code_id = damage['damage_code']
                            existing_damage.damage_cause_code_id = damage.get('damage_cause_code')
                            existing_damage.comments = damage.get('comments')
                            existing_damage.save()
                        else:
                            # Create new record
                            TrtDamage.objects.create(
                                observation=observation,
                                body_part_id=body_part_id,
                                damage_code_id=damage['damage_code'],
                                damage_cause_code_id=damage.get('damage_cause_code'),
                                comments=damage.get('comments')
                            )
                    except IntegrityError as e:
                        print(f"Integrity error processing damage records: {str(e)}")
                        print(f"observation_id={observation.observation_id}, body_part={body_part_id}")
                        continue
                    except Exception as e:
                        print(f"Error processing damage records: {str(e)}")
                        print(traceback.format_exc())
                        continue

        except Exception as e:
            print(f"Error updating damage records: {str(e)}")
            print(traceback.format_exc())
            raise ValidationError(f"Error updating damage records: {str(e)}")

    def _update_scars(self, observation, scars_data):
        """Update scar records"""
        try:
            if scars_data:
                # Update left scars
                observation.scars_left = scars_data.get('scars_left', False)
                observation.scars_left_scale_1 = scars_data.get('scars_left_scale_1', False)
                observation.scars_left_scale_2 = scars_data.get('scars_left_scale_2', False)
                observation.scars_left_scale_3 = scars_data.get('scars_left_scale_3', False)
                
                # Update right scars
                observation.scars_right = scars_data.get('scars_right', False)
                observation.scars_right_scale_1 = scars_data.get('scars_right_scale_1', False)
                observation.scars_right_scale_2 = scars_data.get('scars_right_scale_2', False)
                observation.scars_right_scale_3 = scars_data.get('scars_right_scale_3', False)
                
                # Update check status
                observation.tagscarnotchecked = scars_data.get('tag_scar_not_checked', False)
                observation.didnotcheckforinjury = scars_data.get('did_not_check_for_injury', False)
                
                observation.save()
        except Exception as e:
            print(f"Error updating scar records: {str(e)}")
            print(traceback.format_exc())
            raise ValidationError(f"Error updating scar records: {str(e)}")

    def _update_other_tags(self, observation, other_tags_data):
        """Update other tag information"""
        try:
            if other_tags_data:
                observation.other_tags = other_tags_data.get('other_tags')
                
                # Update identification type
                identification_type = other_tags_data.get('identification_type')
                if identification_type:
                    try:
                        identification_type_obj = TrtIdentificationTypes.objects.get(
                            identification_type=identification_type
                        )
                        observation.other_tags_identification_type = identification_type_obj
                    except TrtIdentificationTypes.DoesNotExist:
                        observation.other_tags_identification_type = None
                else:
                    observation.other_tags_identification_type = None
                    
                observation.save()
        except Exception as e:
            print(f"Error updating other tag information: {str(e)}")
            print(traceback.format_exc())
            raise ValidationError(f"Error updating other tag information: {str(e)}")

    def _update_status(self, observation):
        """Update observation status"""
        try:
            # Calculate status based on business rules
            if observation.nesting and observation.nesting.code == 'Y':
                # Check if it's the first observation
                previous_observations = TrtObservations.objects.filter(
                    turtle_id=observation.turtle_id,
                    observation_date__lt=observation.observation_date
                ).exists()
                
                observation.observation_status = 'S' if previous_observations else 'I'
            else:
                observation.observation_status = 'N'
                
            observation.save()
        except Exception as e:
            print(f"Error updating observation status: {str(e)}")
            print(traceback.format_exc())
            raise ValidationError(f"Error updating observation status: {str(e)}")

    def _validate_data(self, data):
        """Validate input data"""
        try:
            # Validate required fields
            required_fields = ['observation_date', 'turtle_id']
            for field in required_fields:
                if not data.get('basic_info', {}).get(field):
                    raise ValidationError(f"Missing required field: {field}")
            
            # Validate date format
            try:
                datetime.strptime(
                    data['basic_info']['observation_date'], 
                    '%Y-%m-%dT%H:%M'
                )
            except (ValueError, KeyError):
                raise ValidationError("Invalid date format")
            
            # Validate coordinates
            location = data.get('location', {})
            if location.get('latitude') is not None:
                try:
                    lat = float(location['latitude'])
                    if not -90 <= lat <= 90:
                        raise ValidationError("Latitude must be between -90 and 90")
                except ValueError:
                    raise ValidationError("Latitude must be a valid number")
                    
            if location.get('longitude') is not None:
                try:
                    lon = float(location['longitude'])
                    if not -180 <= lon <= 180:
                        raise ValidationError("Longitude must be between -180 and 180")
                except ValueError:
                    raise ValidationError("Longitude must be a valid number")
                    
        except Exception as e:
            print(f"Error validating data: {str(e)}")
            print(traceback.format_exc())
            raise ValidationError(f"Error validating data: {str(e)}")


class TurtleManagementView(TemplateView):
    template_name = 'wamtram2/turtle_management.html'
    
    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        context['sex_choices'] = SEX_CHOICES
        context['cause_of_death_choices'] = TrtCauseOfDeath.objects.all()
        context['turtle_status_choices'] = TrtTurtleStatus.objects.all()
        context['species_choices'] = TrtSpecies.objects.all()
        context['location_choices'] = TrtLocations.objects.all()
        
        return context

    def get(self, request, *args, **kwargs):
        if request.headers.get('X-Requested-With') == 'XMLHttpRequest':
            return self.handle_ajax_request(request)
        return super().get(request, *args, **kwargs)

    def post(self, request, *args, **kwargs):
        if request.headers.get('X-Requested-With') == 'XMLHttpRequest':
            try:
                data = json.loads(request.body)
                
                turtle = TrtTurtles.objects.get(turtle_id=data['turtle_id'])
                
                if data.get('species'):
                    species = TrtSpecies.objects.get(species_code=data['species'])
                    turtle.species_code = species
                
                if data.get('location'):
                    location = TrtLocations.objects.get(location_code=data['location'])
                    turtle.location = location
                    
                if data.get('turtle_status'):
                    status = TrtTurtleStatus.objects.get(turtle_status=data['turtle_status'])
                    turtle.turtle_status = status
                    
                if data.get('cause_of_death'):
                    cod = TrtCauseOfDeath.objects.get(cause_of_death=data['cause_of_death'])
                    turtle.cause_of_death = cod
                    
                turtle.turtle_name = data.get('turtle_name', '')
                turtle.sex = data.get('sex', '')
                turtle.comments = data.get('comments', '')
                
                turtle.save()
                
                return JsonResponse({
                    'status': 'success',
                    'message': 'Turtle information updated successfully'
                })
            except json.JSONDecodeError as e:
                return JsonResponse({
                    'status': 'error',
                    'message': 'Invalid JSON data'
                }, status=400)
            except Exception as e:
                return JsonResponse({
                    'status': 'error',
                    'message': str(e)
                }, status=500)
        return super().post(request, *args, **kwargs)



    def handle_ajax_request(self, request):
        turtle_id = request.GET.get('turtle_id')
        tag_id = request.GET.get('tag_id')
        pit_tag_id = request.GET.get('pit_tag_id')
        other_id = request.GET.get('other_id')

        queryset = None

        if turtle_id:
            queryset = TrtTurtles.objects.filter(turtle_id=turtle_id)
        elif tag_id:
            turtle_ids = TrtTags.objects.filter(
                tag_id__icontains=tag_id
            ).values_list('turtle_id', flat=True)
            queryset = TrtTurtles.objects.filter(turtle_id__in=turtle_ids)
        elif pit_tag_id:
            turtle_ids = TrtPitTags.objects.filter(
                pittag_id__icontains=pit_tag_id
            ).values_list('turtle_id', flat=True)
            queryset = TrtTurtles.objects.filter(turtle_id__in=turtle_ids)
        elif other_id:
            turtle_ids = TrtIdentification.objects.filter(
                identifier__icontains=other_id
            ).values_list('turtle_id', flat=True)
            queryset = TrtTurtles.objects.filter(turtle_id__in=turtle_ids)

        if queryset is None:
            return JsonResponse({
                'status': 'error',
                'message': 'No search criteria provided'
            })

        turtle_data = []
        for turtle in queryset:
            tags = TrtTags.objects.filter(turtle=turtle.turtle_id)
            tag_data = [{
                'tag_id': tag.tag_id,
                'side': tag.side,
                'tag_status': tag.tag_status.description if tag.tag_status else None,
                'issue_location': tag.issue_location,
                'custodian_person': str(tag.custodian_person) if tag.custodian_person else None,
                'return_date': tag.return_date.strftime('%Y-%m-%d') if tag.return_date else None,
                'return_condition': tag.return_condition,
                'comments': tag.comments,
                'field_person_id': tag.field_person_id
            } for tag in tags]
        
            
            pit_tags = TrtPitTags.objects.filter(turtle=turtle.turtle_id)
            pit_tag_data = [{
                'pit_tag_id': tag.pittag_id,
                'issue_location': tag.issue_location,
                'custodian_person': str(tag.custodian_person) if tag.custodian_person else None,
                'pit_tag_status': tag.pit_tag_status.pit_tag_status if tag.pit_tag_status else None,
                'return_date': tag.return_date.strftime('%Y-%m-%d') if tag.return_date else None,
                'return_condition': tag.return_condition,
                'comments': tag.comments,
                'field_person_id': tag.field_person_id,
                'tag_order_id': tag.tag_order_id,
                'batch_number': tag.batch_number,
                'box_number': tag.box_number
            } for tag in pit_tags]
        
            identifications = TrtIdentification.objects.filter(turtle=turtle.turtle_id)
            identification_data = [{
                'identification_type': ident.identification_type.identification_type,
                'identifier': ident.identifier,
                'comments': ident.comments
            } for ident in identifications]
            
            observations = turtle.trtobservations_set.all()
            observation_data = [{
                'observation_id': obs.pk,
                'date_time': obs.observation_date.strftime('%Y-%m-%dT%H:%M'),
                'observation_status': obs.observation_status,
                'alive': str(obs.alive),
                'place': obs.place_code.get_full_name() if obs.place_code else '',
                'nesting': str(obs.nesting),
                'activity': str(obs.activity_code)
            } for obs in observations]
            
            samples = turtle.trtsamples_set.all()
            sample_data = [{
                'tissue_type': sample.tissue_type.tissue_type,
                'observation_id': sample.observation_id,
                'sample_date': sample.sample_date.strftime('%Y-%m-%d') if sample.sample_date else '',
                'label': sample.sample_label,
                'comments': sample.comments
            } for sample in samples]
            
            documents = TrtDocuments.objects.filter(turtle_id=turtle.turtle_id)
            document_data = [{
                'document_id': doc.pk,
                'document_type': str(doc.document_type),
                'file_name': doc.filename,
                'person_id': str(doc.person_id)
            } for doc in documents]
                
            turtle_data.append({
                'turtle_id': turtle.turtle_id,
                'species': turtle.species_code.species_code,
                'turtle_name': turtle.turtle_name or '',
                'sex': turtle.sex if turtle.sex else '',
                'cause_of_death': turtle.cause_of_death.cause_of_death if turtle.cause_of_death else '',
                'turtle_status': turtle.turtle_status.turtle_status if turtle.turtle_status else '',
                'date_entered': turtle.date_entered.strftime('%Y-%m-%d') if turtle.date_entered else '',
                'comments': turtle.comments or '',
                'location': turtle.location_code.location_code if turtle.location_code else '',
                'tags': tag_data,
                'pit_tags': pit_tag_data,
                'identifications': identification_data,
                'observations': observation_data,
                'samples': sample_data,
                'documents': document_data
            })
            

        return JsonResponse({
            'status': 'success',
            'data': turtle_data
        })
        
        